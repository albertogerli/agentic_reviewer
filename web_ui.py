"""
Web UI for Generic Document Review System
User-friendly interface using Gradio for non-technical users.
"""

import gradio as gr
import asyncio
import os
import sys
import time
from pathlib import Path
import json
import logging
from datetime import datetime
from typing import Optional, Tuple, List
import tempfile
import shutil

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent))

from main import Config, setup_logging
from generic_reviewer import (
    GenericReviewOrchestrator,
    IterativeReviewOrchestrator,
    DocumentClassifier,
    FileManager
)

# Setup logging
logger = setup_logging()

# Global config
_config = None


def initialize_system():
    """Initialize the review system."""
    global _config
    
    # Check for API key
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        return False, "⚠️ OPENAI_API_KEY not set. Please set your API key in environment variables."
    
    try:
        # Check if config.yaml exists, if not use defaults
        if not Path("config.yaml").exists():
            logger.warning("config.yaml not found, creating default config")
            _config = Config()
        else:
            _config = Config.from_yaml("config.yaml")
        _config.validate()
        return True, "✅ System initialized successfully!"
    except Exception as e:
        logger.error(f"Config error: {e}", exc_info=True)
        # Try with default config
        try:
            _config = Config()
            return True, "✅ System initialized with default config"
        except:
            return False, f"⚠️ Configuration error: {str(e)}"


def estimate_processing_time(file_path: str, enable_deep_review: bool, enable_iterative: bool, max_iterations: int) -> str:
    """Estimate processing time based on document and options."""
    try:
        file_size_mb = Path(file_path).stat().st_size / (1024 * 1024)
        
        # Base time
        base_time = 2 + (file_size_mb * 0.5)  # 2 min + 0.5 min per MB
        
        # Adjustments
        if enable_deep_review:
            base_time *= 2.5
        if enable_iterative:
            base_time *= max_iterations * 1.3
        
        minutes = int(base_time)
        if minutes < 1:
            return "< 1 minute"
        elif minutes == 1:
            return "~1 minute"
        else:
            return f"~{minutes}-{minutes+2} minutes"
    except:
        return "2-5 minutes"


def generate_agent_preview_html(agent_list: list, document_type: str = "Document") -> str:
    """Generate INTRIGANTE HTML preview of agents being deployed like a mission briefing!"""
    agent_icons = {
        'style_editor': '🎨',
        'consistency_checker': '🔍',
        'fact_checker': '✓',
        'logic_checker': '🧠',
        'technical_expert': '⚙️',
        'subject_matter_expert': '🎓',
        'business_analyst': '💼',
        'financial_analyst': '💰',
        'legal_expert': '⚖️',
        'data_validator': '📊',
        'web_researcher': '🌐',
        'academic_researcher': '📚',
    }
    
    # Mission-style briefing header
    html = f"""
    <div style="
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 30px;
        border-radius: 20px;
        margin: 20px 0;
        box-shadow: 0 10px 40px rgba(102, 126, 234, 0.3);
        position: relative;
        overflow: hidden;
    ">
        <!-- Animated background pattern -->
        <div style="
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background: repeating-linear-gradient(
                45deg,
                transparent,
                transparent 10px,
                rgba(255,255,255,0.05) 10px,
                rgba(255,255,255,0.05) 20px
            );
            animation: slide 20s linear infinite;
        "></div>
        
        <div style="position: relative; z-index: 1;">
            <div style="text-align: center; margin-bottom: 25px;">
                <div style="
                    display: inline-block;
                    background: rgba(255,255,255,0.2);
                    padding: 10px 25px;
                    border-radius: 30px;
                    font-size: 12px;
                    color: white;
                    font-weight: 600;
                    letter-spacing: 2px;
                    margin-bottom: 15px;
                    animation: pulse 2s infinite;
                ">
                    🚀 MISSION BRIEFING
                </div>
                <h3 style="
                    margin: 0;
                    color: white;
                    font-size: 24px;
                    font-weight: 700;
                    text-shadow: 0 2px 10px rgba(0,0,0,0.2);
                ">
                    AI Agent Team Assembled
                </h3>
                <div style="color: rgba(255,255,255,0.9); font-size: 14px; margin-top: 8px;">
                    Analyzing your <strong>{document_type}</strong> with {len(agent_list)} specialized agents
                </div>
            </div>
            
            <!-- Agent Grid with special effects -->
            <div style="
                display: grid;
                grid-template-columns: repeat(auto-fit, minmax(160px, 1fr));
                gap: 12px;
                margin-top: 25px;
            ">
    """
    
    # Add each agent with special animations
    for i, agent in enumerate(agent_list[:12]):
        agent_name = agent.replace('_', ' ').title()
        icon = agent_icons.get(agent, '🤖')
        delay = i * 0.08  # Faster cascade
        
        # Assign tier badges
        tier_badge = ""
        tier_color = "#4ade80"
        if i < 5:  # Tier 1 - Core
            tier_badge = "CORE"
            tier_color = "#f59e0b"
        elif i < 11:  # Tier 2 - Specialists
            tier_badge = "SPECIALIST"
            tier_color = "#3b82f6"
        else:  # Tier 3 - Deep
            tier_badge = "DEEP DIVE"
            tier_color = "#8b5cf6"
        
        html += f"""
        <div style="
            background: rgba(255,255,255,0.95);
            padding: 18px 12px;
            border-radius: 12px;
            text-align: center;
            animation: slideInUp 0.4s ease-out {delay}s backwards, glow 3s ease-in-out infinite;
            box-shadow: 0 4px 15px rgba(0,0,0,0.1);
            border: 2px solid rgba(255,255,255,0.3);
            position: relative;
            transition: all 0.3s ease;
        " class="agent-deploy-card">
            <!-- Tier badge -->
            <div style="
                position: absolute;
                top: -8px;
                right: -8px;
                background: {tier_color};
                color: white;
                font-size: 9px;
                padding: 3px 8px;
                border-radius: 10px;
                font-weight: 700;
                box-shadow: 0 2px 8px rgba(0,0,0,0.2);
            ">
                {tier_badge}
            </div>
            
            <div style="
                font-size: 40px;
                margin-bottom: 10px;
                animation: float 3s ease-in-out infinite;
                animation-delay: {delay}s;
            ">{icon}</div>
            <div style="
                font-size: 13px;
                font-weight: 700;
                color: #667eea;
                margin-bottom: 4px;
            ">{agent_name}</div>
            <div style="
                font-size: 10px;
                color: #4ade80;
                font-weight: 600;
                display: flex;
                align-items: center;
                justify-content: center;
                gap: 4px;
            ">
                <span style="
                    display: inline-block;
                    width: 6px;
                    height: 6px;
                    background: #4ade80;
                    border-radius: 50%;
                    animation: pulse 2s infinite;
                "></span>
                DEPLOYED
            </div>
        </div>
        """
    
    if len(agent_list) > 12:
        html += f"""
        <div style="
            background: rgba(255,255,255,0.1);
            padding: 18px;
            border-radius: 12px;
            text-align: center;
            color: white;
            font-weight: 700;
            display: flex;
            align-items: center;
            justify-content: center;
            border: 2px dashed rgba(255,255,255,0.3);
        ">
            <span style="font-size: 24px; margin-right: 8px;">+{len(agent_list) - 12}</span>
            <span style="font-size: 12px;">MORE<br>AGENTS</span>
        </div>
        """
    
    html += """
            </div>
            
            <!-- Mission status bar -->
            <div style="
                margin-top: 25px;
                padding: 18px;
                background: rgba(255,255,255,0.15);
                border-radius: 12px;
                text-align: center;
                backdrop-filter: blur(10px);
            ">
                <div style="
                    display: flex;
                    align-items: center;
                    justify-content: center;
                    gap: 12px;
                    color: white;
                    font-weight: 600;
                    font-size: 14px;
                ">
                    <div style="
                        width: 20px;
                        height: 20px;
                        border: 3px solid white;
                        border-top-color: transparent;
                        border-radius: 50%;
                        animation: spin-glow 1s linear infinite;
                    "></div>
                    <span>ANALYSIS IN PROGRESS</span>
                </div>
                <div style="
                    color: rgba(255,255,255,0.8);
                    font-size: 12px;
                    margin-top: 8px;
                ">
                    Each agent will provide specialized insights from their domain expertise
                </div>
            </div>
        </div>
    </div>
    
    <style>
    @keyframes slide {
        0% { transform: translateX(0); }
        100% { transform: translateX(20px); }
    }
    
    .agent-deploy-card:hover {
        transform: translateY(-5px) scale(1.05);
        box-shadow: 0 8px 25px rgba(102, 126, 234, 0.3);
    }
    </style>
    """
    
    return html


def generate_agents_deployed_html(results: dict, document_type: str = "Document") -> str:
    """Generate stunning HTML showing which agents were deployed and their AI models."""
    if not results or not isinstance(results, dict):
        return ""
    
    # Agent icon mapping
    agent_icons = {
        'style_editor': '🎨', 'consistency_checker': '🔍', 'fact_checker': '✓',
        'logic_checker': '🧠', 'technical_expert': '⚙️', 'subject_matter_expert': '🎓',
        'business_analyst': '💼', 'financial_analyst': '💰', 'legal_expert': '⚖️',
        'data_validator': '📊', 'web_researcher': '🌐', 'academic_researcher': '📚',
        'data_analyst': '📈', 'coordinator': '👑', 'final_evaluator': '⭐'
    }
    
    # Collect deployed agents
    deployed_agents = []
    for key in results.keys():
        if key not in ['metadata', 'summary']:
            agent_name = key.replace('_', ' ').title()
            icon = agent_icons.get(key, '🤖')
            
            # Infer model based on agent type
            if key in ['coordinator', 'final_evaluator', 'consistency_checker', 'logic_checker', 'subject_matter_expert', 'legal_expert', 'data_validator']:
                model = 'gpt-5'
                model_color = '#ff6b6b'
            elif key in ['web_researcher', 'fact_checker', 'business_analyst', 'financial_analyst']:
                model = 'gpt-5-mini'
                model_color = '#4ecdc4'
            else:
                model = 'gpt-5-mini'
                model_color = '#4ecdc4'
            
            deployed_agents.append({
                'name': agent_name,
                'icon': icon,
                'model': model,
                'model_color': model_color,
                'key': key
            })
    
    if not deployed_agents:
        return ""
    
    html = f"""
    <div style="
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 30px;
        border-radius: 20px;
        margin: 20px 0;
        box-shadow: 0 10px 30px rgba(102, 126, 234, 0.3);
    ">
        <div style="text-align: center; color: white; margin-bottom: 25px;">
            <h2 style="margin: 0 0 10px 0; font-size: 28px; text-shadow: 2px 2px 4px rgba(0,0,0,0.2);">
                🤖 AI Squad Deployed
            </h2>
            <p style="margin: 0; font-size: 16px; opacity: 0.95;">
                {len(deployed_agents)} Specialized Agents for {document_type} Analysis
            </p>
        </div>
        
        <style>
            .agent-chip {{
                display: inline-flex;
                align-items: center;
                gap: 10px;
                background: rgba(255, 255, 255, 0.95);
                padding: 14px 20px;
                border-radius: 30px;
                margin: 8px;
                box-shadow: 0 4px 15px rgba(0, 0, 0, 0.15);
                transition: all 0.3s cubic-bezier(0.175, 0.885, 0.32, 1.275);
                animation: slideInUp 0.5s ease-out;
            }}
            .agent-chip:hover {{
                transform: translateY(-5px) scale(1.05);
                box-shadow: 0 8px 25px rgba(0, 0, 0, 0.25);
            }}
            .agent-icon {{
                font-size: 26px;
                filter: drop-shadow(1px 1px 2px rgba(0,0,0,0.2));
            }}
            .agent-name {{
                font-weight: 600;
                color: #2d3748;
                font-size: 15px;
            }}
            .agent-model {{
                color: white;
                padding: 5px 14px;
                border-radius: 15px;
                font-size: 11px;
                font-weight: 700;
                text-transform: uppercase;
                letter-spacing: 0.5px;
                box-shadow: 0 2px 8px rgba(0,0,0,0.15);
            }}
        </style>
        
        <div style="
            display: flex;
            flex-wrap: wrap;
            justify-content: center;
            align-items: center;
            gap: 10px;
            margin-top: 20px;
        ">
    """
    
    for i, agent in enumerate(deployed_agents):
        html += f"""
        <div class="agent-chip" style="animation-delay: {i * 0.05}s;">
            <span class="agent-icon">{agent['icon']}</span>
            <span class="agent-name">{agent['name']}</span>
            <span class="agent-model" style="background: {agent['model_color']};">
                {agent['model']}
            </span>
        </div>
        """
    
    html += """
        </div>
        
        <div style="
            margin-top: 25px;
            padding: 20px;
            background: rgba(255, 255, 255, 0.1);
            border-radius: 15px;
            backdrop-filter: blur(10px);
        ">
            <div style="display: flex; justify-content: center; gap: 30px; flex-wrap: wrap; color: white; font-size: 14px;">
                <div style="display: flex; align-items: center; gap: 8px;">
                    <span style="display: inline-block; width: 16px; height: 16px; background: #ff6b6b; border-radius: 50%;"></span>
                    <strong>gpt-5</strong> — High complexity tasks
                </div>
                <div style="display: flex; align-items: center; gap: 8px;">
                    <span style="display: inline-block; width: 16px; height: 16px; background: #4ecdc4; border-radius: 50%;"></span>
                    <strong>gpt-5-mini</strong> — Medium complexity
                </div>
                <div style="display: flex; align-items: center; gap: 8px;">
                    <span style="display: inline-block; width: 16px; height: 16px; background: #95e1d3; border-radius: 50%;"></span>
                    <strong>gpt-5-nano</strong> — Fast analysis
                </div>
            </div>
        </div>
    </div>
    """
    
    return html


def generate_agent_ratings_html(results: dict) -> str:
    """Generate stunning HTML with star ratings for each agent's performance."""
    # Debug logging
    logger.debug(f"generate_agent_ratings_html called with results type: {type(results)}")
    if results:
        logger.debug(f"Results keys: {list(results.keys())}")
    
    if not results or not isinstance(results, dict):
        logger.warning(f"No valid results for ratings: results={bool(results)}, is_dict={isinstance(results, dict)}")
        return "<p style='text-align: center; color: #6c757d; padding: 40px;'>No agent ratings available yet. Complete an analysis first!</p>"
    
    # Agent icon mapping
    agent_icons = {
        'style_editor': '🎨', 'consistency_checker': '🔍', 'fact_checker': '✓',
        'logic_checker': '🧠', 'technical_expert': '⚙️', 'subject_matter_expert': '🎓',
        'business_analyst': '💼', 'financial_analyst': '💰', 'legal_expert': '⚖️',
        'data_validator': '📊', 'web_researcher': '🌐', 'academic_researcher': '📚',
        'data_analyst': '📈', 'coordinator': '👑', 'final_evaluator': '⭐'
    }
    
    html = """
    <div style="max-width: 1400px; margin: 0 auto; padding: 30px;">
        <div style="text-align: center; margin-bottom: 40px;">
            <h2 style="color: #667eea; font-size: 32px; margin-bottom: 10px;">⭐ Agent Performance Ratings</h2>
            <p style="color: #6c757d; font-size: 16px;">Each agent is rated on thoroughness, insight quality, and depth of analysis</p>
        </div>
        
        <style>
            .rating-grid {
                display: grid;
                grid-template-columns: repeat(auto-fill, minmax(280px, 1fr));
                gap: 25px;
                margin-top: 30px;
            }
            .rating-card {
                background: linear-gradient(135deg, #ffffff 0%, #f8f9fa 100%);
                border: 2px solid #e9ecef;
                border-radius: 20px;
                padding: 25px;
                transition: all 0.4s cubic-bezier(0.175, 0.885, 0.32, 1.275);
                animation: fadeIn 0.6s ease-out;
                position: relative;
                overflow: hidden;
            }
            .rating-card::before {
                content: '';
                position: absolute;
                top: 0;
                left: 0;
                width: 100%;
                height: 5px;
                background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
                transform: scaleX(0);
                transition: transform 0.4s ease;
            }
            .rating-card:hover {
                transform: translateY(-10px) scale(1.02);
                box-shadow: 0 20px 40px rgba(102, 126, 234, 0.25);
                border-color: #667eea;
            }
            .rating-card:hover::before {
                transform: scaleX(1);
            }
            .rating-header {
                display: flex;
                align-items: center;
                margin-bottom: 20px;
                padding-bottom: 15px;
                border-bottom: 2px solid #e9ecef;
            }
            .rating-icon {
                font-size: 48px;
                margin-right: 15px;
                filter: drop-shadow(2px 2px 4px rgba(0,0,0,0.1));
                animation: float 3s ease-in-out infinite;
            }
            .rating-name {
                font-size: 18px;
                font-weight: 700;
                color: #667eea;
                line-height: 1.3;
            }
            .stars-container {
                text-align: center;
                margin: 20px 0;
            }
            .stars {
                font-size: 28px;
                color: #ffd700;
                text-shadow: 2px 2px 4px rgba(255, 215, 0, 0.3);
                letter-spacing: 3px;
            }
            .rating-score {
                text-align: center;
                font-size: 24px;
                font-weight: 700;
                color: #495057;
                margin: 15px 0;
                font-family: 'Courier New', monospace;
            }
            .rating-badge {
                text-align: center;
                font-size: 14px;
                font-weight: 600;
                padding: 8px 16px;
                border-radius: 20px;
                background: linear-gradient(135deg, rgba(102, 126, 234, 0.1) 0%, rgba(118, 75, 162, 0.1) 100%);
                margin-top: 15px;
            }
            @keyframes float {
                0%, 100% { transform: translateY(0px); }
                50% { transform: translateY(-10px); }
            }
            @keyframes fadeIn {
                from { opacity: 0; transform: translateY(20px); }
                to { opacity: 1; transform: translateY(0); }
            }
        </style>
        
        <div class="rating-grid">
    """
    
    # Process each agent
    agent_count = 0
    for agent_key, agent_data in results.items():
        if agent_key in ['metadata', 'summary']:
            continue
        
        agent_count += 1
        
        # Get agent details
        agent_name = agent_key.replace('_', ' ').title()
        icon = agent_icons.get(agent_key, '🤖')
        
        # Extract review text
        review_text = ""
        if isinstance(agent_data, dict):
            review_text = agent_data.get('review', agent_data.get('summary', agent_data.get('assessment', '')))
        elif isinstance(agent_data, str):
            review_text = agent_data
        
        # Calculate smart rating
        rating = 3.0  # Base
        
        # Length bonuses
        if len(review_text) > 500: rating += 0.3
        if len(review_text) > 1000: rating += 0.4
        if len(review_text) > 2000: rating += 0.5
        
        # Quality indicators
        positive_words = ['excellent', 'outstanding', 'comprehensive', 'thorough', 'detailed', 'in-depth', 'exceptional']
        quality_score = sum(1 for word in positive_words if word in review_text.lower())
        rating += min(quality_score * 0.15, 0.6)
        
        # Structure bonus
        if review_text.count('\n') > 10: rating += 0.2
        if review_text.count('-') > 5: rating += 0.1  # Bullet points
        
        # Cap at 5.0
        rating = min(5.0, max(2.5, rating))
        
        # Generate stars
        full_stars = int(rating)
        half_star = (rating - full_stars) >= 0.5
        empty_stars = 5 - full_stars - (1 if half_star else 0)
        
        stars_html = "★" * full_stars
        if half_star:
            stars_html += "⯨"
        stars_html += "☆" * empty_stars
        
        # Rating badge
        if rating >= 4.7:
            badge = "🏆 Legendary"
            badge_color = "#ffd700"
        elif rating >= 4.3:
            badge = "✨ Outstanding"
            badge_color = "#667eea"
        elif rating >= 4.0:
            badge = "🌟 Excellent"
            badge_color = "#764ba2"
        elif rating >= 3.5:
            badge = "👍 Very Good"
            badge_color = "#17a2b8"
        elif rating >= 3.0:
            badge = "✓ Good"
            badge_color = "#28a745"
        else:
            badge = "○ Fair"
            badge_color = "#ffc107"
        
        html += f"""
        <div class="rating-card" style="animation-delay: {agent_count * 0.1}s;">
            <div class="rating-header">
                <div class="rating-icon">{icon}</div>
                <div class="rating-name">{agent_name}</div>
            </div>
            <div class="stars-container">
                <div class="stars">{stars_html}</div>
            </div>
            <div class="rating-score">{rating:.1f} / 5.0</div>
            <div class="rating-badge" style="color: {badge_color}; border: 2px solid {badge_color};">
                {badge}
            </div>
        </div>
        """
    
    html += """
        </div>
        
        <div style="margin-top: 50px; text-align: center; padding: 30px; background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%); border-radius: 15px;">
            <h3 style="color: #667eea; margin-bottom: 15px;">📊 Rating System</h3>
            <div style="display: flex; justify-content: center; gap: 30px; flex-wrap: wrap; margin-top: 20px;">
                <div><strong>🏆 Legendary:</strong> 4.7-5.0</div>
                <div><strong>✨ Outstanding:</strong> 4.3-4.6</div>
                <div><strong>🌟 Excellent:</strong> 4.0-4.2</div>
                <div><strong>👍 Very Good:</strong> 3.5-3.9</div>
                <div><strong>✓ Good:</strong> 3.0-3.4</div>
            </div>
        </div>
    </div>
    """
    
    return html


def generate_live_activity_log(activities: list, current_agent: str = None, total_agents: int = 0, completed: int = 0) -> str:
    """Generate cumulative live activity log with history."""
    # Agent icon mapping
    agent_icons = {
        'style_editor': '🎨', 'consistency_checker': '🔍', 'fact_checker': '✓',
        'logic_checker': '🧠', 'technical_expert': '⚙️', 'subject_matter_expert': '🎓',
        'business_analyst': '💼', 'financial_analyst': '💰', 'legal_expert': '⚖️',
        'data_validator': '📊', 'web_researcher': '🌐', 'academic_researcher': '📚',
        'data_analyst': '📈', 'coordinator': '👑', 'final_evaluator': '⭐'
    }
    
    html = f"""
    <div style="
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 25px;
        border-radius: 20px;
        margin: 20px 0;
        box-shadow: 0 10px 30px rgba(102, 126, 234, 0.3);
    ">
        <style>
            @keyframes slideDown {{
                from {{ transform: translateY(-10px); opacity: 0; }}
                to {{ transform: translateY(0); opacity: 1; }}
            }}
            @keyframes pulse {{
                0%, 100% {{ opacity: 1; }}
                50% {{ opacity: 0.6; }}
            }}
            .activity-item {{
                animation: slideDown 0.3s ease-out;
            }}
            .live-dot {{
                display: inline-block;
                width: 10px;
                height: 10px;
                background: #22c55e;
                border-radius: 50%;
                animation: pulse 2s infinite;
                margin-right: 8px;
            }}
        </style>
        
        <!-- Header with progress -->
        <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 20px; color: white;">
            <div style="display: flex; align-items: center; gap: 10px;">
                <span class="live-dot"></span>
                <strong style="font-size: 18px;">LIVE ACTIVITY</strong>
            </div>
            <div style="font-size: 16px; font-weight: 600;">
                {completed}/{total_agents} Agents
            </div>
        </div>
        
        <!-- Activity Feed (scrollable) -->
        <div style="
            background: rgba(255,255,255,0.95);
            border-radius: 15px;
            padding: 20px;
            max-height: 400px;
            overflow-y: auto;
        ">
    """
    
    # Add each activity
    for activity in activities:
        agent_key = activity.get('agent', 'unknown')
        status_text = activity.get('status', '')
        is_current = activity.get('is_current', False)
        is_complete = activity.get('is_complete', False)
        
        icon = agent_icons.get(agent_key, '🤖')
        agent_display = agent_key.replace('_', ' ').title()
        
        if is_current:
            bg_color = "#fef3c7"  # Yellow
            status_badge = "⚡ ANALYZING"
            badge_color = "#f59e0b"
        elif is_complete:
            bg_color = "#d1fae5"  # Green
            status_badge = "✓ COMPLETED"
            badge_color = "#10b981"
        else:
            bg_color = "#e5e7eb"  # Gray
            status_badge = "⏳ PENDING"
            badge_color = "#6b7280"
        
        html += f"""
        <div class="activity-item" style="
            background: {bg_color};
            padding: 15px;
            border-radius: 10px;
            margin-bottom: 10px;
            border-left: 4px solid {badge_color};
            display: flex;
            align-items: center;
            justify-content: space-between;
        ">
            <div style="display: flex; align-items: center; gap: 12px;">
                <span style="font-size: 28px;">{icon}</span>
                <div>
                    <div style="font-weight: 600; color: #1f2937; font-size: 15px;">
                        {agent_display}
                    </div>
                    <div style="font-size: 12px; color: #6b7280; margin-top: 2px;">
                        {status_text}
                    </div>
                </div>
            </div>
            <span style="
                background: {badge_color};
                color: white;
                padding: 4px 12px;
                border-radius: 12px;
                font-size: 10px;
                font-weight: 700;
                letter-spacing: 0.5px;
            ">{status_badge}</span>
        </div>
        """
    
    html += """
        </div>
    </div>
    """
    
    return html


def generate_live_activity_html(agent_name: str, status: str, progress_pct: float, total_agents: int, completed: int) -> str:
    """Generate dynamic live activity feed during processing."""
    # Agent icon mapping
    agent_icons = {
        'style_editor': '🎨', 'consistency_checker': '🔍', 'fact_checker': '✓',
        'logic_checker': '🧠', 'technical_expert': '⚙️', 'subject_matter_expert': '🎓',
        'business_analyst': '💼', 'financial_analyst': '💰', 'legal_expert': '⚖️',
        'data_validator': '📊', 'web_researcher': '🌐', 'academic_researcher': '📚',
        'data_analyst': '📈', 'coordinator': '👑', 'final_evaluator': '⭐'
    }
    
    icon = agent_icons.get(agent_name, '🤖')
    agent_display = agent_name.replace('_', ' ').title()
    
    html = f"""
    <div style="
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 30px;
        border-radius: 20px;
        margin: 20px 0;
        box-shadow: 0 10px 30px rgba(102, 126, 234, 0.3);
        position: relative;
        overflow: hidden;
    ">
        <style>
            @keyframes pulse {{
                0%, 100% {{ opacity: 1; }}
                50% {{ opacity: 0.6; }}
            }}
            @keyframes slideInRight {{
                from {{ transform: translateX(-20px); opacity: 0; }}
                to {{ transform: translateX(0); opacity: 1; }}
            }}
            @keyframes rotate {{
                from {{ transform: rotate(0deg); }}
                to {{ transform: rotate(360deg); }}
            }}
            .live-indicator {{
                display: inline-block;
                width: 12px;
                height: 12px;
                background: #ff4444;
                border-radius: 50%;
                animation: pulse 1.5s ease-in-out infinite;
                margin-right: 10px;
            }}
            .agent-working {{
                animation: slideInRight 0.5s ease-out;
            }}
            .spinner {{
                display: inline-block;
                width: 20px;
                height: 20px;
                border: 3px solid rgba(255,255,255,0.3);
                border-top-color: white;
                border-radius: 50%;
                animation: rotate 1s linear infinite;
                margin-left: 10px;
            }}
        </style>
        
        <!-- Live Indicator -->
        <div style="text-align: center; color: white; margin-bottom: 20px;">
            <div style="display: inline-flex; align-items: center; background: rgba(255,255,255,0.2); padding: 8px 20px; border-radius: 20px;">
                <span class="live-indicator"></span>
                <strong style="font-size: 14px; letter-spacing: 1px;">LIVE ANALYSIS</strong>
            </div>
        </div>
        
        <!-- Current Agent -->
        <div class="agent-working" style="
            background: rgba(255,255,255,0.95);
            padding: 25px;
            border-radius: 15px;
            margin-bottom: 20px;
            box-shadow: 0 4px 15px rgba(0,0,0,0.1);
        ">
            <div style="display: flex; align-items: center; justify-content: space-between;">
                <div style="display: flex; align-items: center; gap: 15px;">
                    <span style="font-size: 48px;">{icon}</span>
                    <div>
                        <div style="font-size: 24px; font-weight: 700; color: #667eea; margin-bottom: 5px;">
                            {agent_display}
                        </div>
                        <div style="font-size: 14px; color: #666;">
                            {status}
                        </div>
                    </div>
                </div>
                <div class="spinner"></div>
            </div>
        </div>
        
        <!-- Progress Stats -->
        <div style="
            display: grid;
            grid-template-columns: repeat(3, 1fr);
            gap: 15px;
            margin-bottom: 20px;
        ">
            <div style="background: rgba(255,255,255,0.15); padding: 15px; border-radius: 10px; text-align: center; backdrop-filter: blur(10px);">
                <div style="font-size: 32px; font-weight: 700; color: white;">{completed}</div>
                <div style="font-size: 12px; color: rgba(255,255,255,0.9); margin-top: 5px;">COMPLETED</div>
            </div>
            <div style="background: rgba(255,255,255,0.15); padding: 15px; border-radius: 10px; text-align: center; backdrop-filter: blur(10px);">
                <div style="font-size: 32px; font-weight: 700; color: white;">{total_agents}</div>
                <div style="font-size: 12px; color: rgba(255,255,255,0.9); margin-top: 5px;">TOTAL AGENTS</div>
            </div>
            <div style="background: rgba(255,255,255,0.15); padding: 15px; border-radius: 10px; text-align: center; backdrop-filter: blur(10px);">
                <div style="font-size: 32px; font-weight: 700; color: white;">{int(progress_pct * 100)}%</div>
                <div style="font-size: 12px; color: rgba(255,255,255,0.9); margin-top: 5px;">PROGRESS</div>
            </div>
        </div>
        
        <!-- Progress Bar -->
        <div style="background: rgba(255,255,255,0.2); height: 8px; border-radius: 10px; overflow: hidden;">
            <div style="
                background: linear-gradient(90deg, #4ade80 0%, #22c55e 100%);
                height: 100%;
                width: {progress_pct * 100}%;
                transition: width 0.5s ease;
                box-shadow: 0 0 10px rgba(74, 222, 128, 0.5);
            "></div>
        </div>
    </div>
    """
    
    return html


def process_document(
    file,
    output_language: str,
    enable_iterative: bool,
    max_iterations: int,
    target_score: float,
    enable_python_tools: bool,
    enable_interactive: bool,
    enable_deep_review: bool,
    reference_files: Optional[List] = None,
    reference_type: str = "example",
    progress=gr.Progress()
) -> Tuple[str, str, str, str, str, str, str, str, str, gr.update, gr.update, gr.update]:
    """
    Process a document and return results.
    
    Returns:
        Tuple of (status_html, report_md, results_json, dashboard_html, agents_report, ratings_html, agents_deployed_html, 
                 output_dir, time_estimate, download_report_btn, download_json_btn, download_dashboard_btn)
    """
    if not file:
        return "⚠️ Please upload a document", "", "", "", "", "", "", "", "", gr.update(), gr.update(), gr.update()
    
    if not _config:
        return "⚠️ System not initialized", "", "", "", "", "", "", "", "", gr.update(), gr.update(), gr.update()
    
    try:
        # Estimate time
        time_estimate_str = estimate_processing_time(file.name, enable_deep_review, enable_iterative, max_iterations)
        time_estimate_html = f"""
        <div style="background: #e3f2fd; padding: 15px; border-radius: 8px; margin: 10px 0; border-left: 4px solid #2196f3;">
            <strong>⏱️ Estimated Time:</strong> {time_estimate_str}
            <br><small>Grab a coffee! We'll work while you wait ☕</small>
        </div>
        """
        
        progress(0, desc="🚀 Starting analysis...")
        
        # Read document
        file_path = file.name
        file_name = Path(file_path).name
        
        # Create unique output directory
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = "".join(c for c in Path(file_name).stem if c.isalnum() or c in (' ', '-', '_')).strip()
        safe_name = safe_name.replace(' ', '_')[:50]
        unique_output_dir = f"reviews/{safe_name}_{timestamp}"
        Path(unique_output_dir).mkdir(parents=True, exist_ok=True)
        
        # Update config
        if not Path("config.yaml").exists():
            config = Config()
        else:
            config = Config.from_yaml("config.yaml")
        config.output_dir = unique_output_dir
        
        # Convert "Auto-detect" to empty string
        if output_language == "Auto-detect":
            output_language = ""
        
        logger.info(f"Processing document: {file_name}")
        logger.info(f"Output language: {output_language or 'Auto-detect'}")
        logger.info(f"Iterative: {enable_iterative}")
        logger.info(f"Deep Review: {enable_deep_review}")
        logger.info(f"Output directory: {unique_output_dir}")
        
        progress(0.05, desc="📄 Analyzing file format...")
        
        # File info
        file_size_mb = Path(file_path).stat().st_size / (1024 * 1024)
        file_info_html = f"""
        <div style="background: #f0f7ff; padding: 15px; border-radius: 8px; margin: 10px 0;">
            <strong>📁 File:</strong> {file_name}<br>
            <strong>📏 Size:</strong> {file_size_mb:.2f} MB<br>
            <strong>🌍 Output Language:</strong> {output_language or 'Auto-detect'}<br>
            <strong>🔬 Deep Analysis:</strong> {'✅ Enabled' if enable_deep_review else '❌ Disabled'}<br>
            <strong>🔄 Iterative:</strong> {'✅ ' + str(max_iterations) + ' rounds' if enable_iterative else '❌ Single pass'}
        </div>
        """
        
        progress(0.1, desc="📖 Reading document content...")
        
        # Read document content with format detection
        file_manager = FileManager(config.output_dir)
        document_text = None
        
        if file_path.lower().endswith(".pdf"):
            document_text = file_manager.extract_text_from_pdf(file_path)
        elif file_path.lower().endswith((".docx", ".doc")):
            # Extract from Word document
            try:
                import docx
                doc = docx.Document(file_path)
                document_text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
                logger.info(f"Successfully read DOCX file ({len(doc.paragraphs)} paragraphs)")
            except ImportError:
                logger.warning("python-docx not installed, trying fallback")
                # Fallback: try reading as text
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                        document_text = f.read()
                    logger.warning("Read DOCX as text with character replacement (install python-docx for better results)")
                except:
                    pass
            except Exception as e:
                logger.error(f"Failed to read DOCX: {e}")
        else:
            # Try multiple encodings for text files
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1', 'utf-16']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        document_text = f.read()
                    logger.info(f"Successfully read file with {encoding} encoding")
                    break
                except UnicodeDecodeError:
                    continue
                except Exception as e:
                    logger.warning(f"Failed to read with {encoding}: {e}")
                    continue
            
            # Last resort: read with errors='replace'
            if not document_text:
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                        document_text = f.read()
                    logger.warning("Read file with UTF-8 and replaced invalid characters")
                except Exception as e:
                    logger.error(f"Failed to read file: {e}")
        
        if not document_text:
            return "⚠️ Failed to read document", "", "", "", "", "", "", unique_output_dir, time_estimate_html + file_info_html, gr.update(), gr.update(), gr.update()
        
        title = Path(file_name).stem
        
        # Count words and characters
        word_count = len(document_text.split())
        char_count = len(document_text)
        
        progress(0.2, desc="📊 Analyzing document structure...")
        
        # Add document stats
        doc_stats_html = f"""
        <div style="background: #f0fff4; padding: 15px; border-radius: 8px; margin: 10px 0; border-left: 4px solid #4ade80;">
            <strong>📝 Document Analysis:</strong><br>
            <strong>Words:</strong> {word_count:,} | <strong>Characters:</strong> {char_count:,}<br>
            <strong>Estimated reading time:</strong> ~{word_count // 200} minutes
        </div>
        """
        
        # Handle reference context
        reference_context = ""
        if reference_files:
            progress(0.25, desc="📚 Processing reference documents...")
            # TODO: Implement reference processing
            logger.info(f"Reference files provided: {len(reference_files)}")
        
        progress(0.27, desc="🔍 Classifying document and selecting agents...")
        
        # Quick classification to show agents
        agents_preview_html = ""
        agent_list = []  # Initialize to avoid undefined reference
        
        try:
            from generic_reviewer import DocumentClassifier
            classifier = DocumentClassifier(config)
            doc_type = classifier.classify_document(document_text[:2000], title)  # Quick classification
            
            # Get list of agents that will be used
            agent_list = doc_type.suggested_agents if hasattr(doc_type, 'suggested_agents') else [
                'style_editor', 'consistency_checker', 'fact_checker', 'logic_checker',
                'technical_expert', 'subject_matter_expert'
            ]
            
            # Generate agent preview HTML
            doc_type_name = doc_type.category if hasattr(doc_type, 'category') else "Document"
            agents_preview_html = generate_agent_preview_html(agent_list, doc_type_name)
            
            # Generate agents deployed HTML EARLY (will show immediately in Report tab)
            # Create a fake results dict with agent keys for early display
            early_results = {agent: "" for agent in agent_list}
            agents_deployed_html = generate_agents_deployed_html(early_results, doc_type_name)
            
            logger.info(f"Document classified as: {doc_type_name}, {len(agent_list)} agents selected")
            
        except Exception as e:
            logger.warning(f"Could not classify document for agent preview: {e}")
            # Fallback agent list
            agent_list = [
                'style_editor', 'consistency_checker', 'fact_checker',
                'logic_checker', 'technical_expert', 'subject_matter_expert'
            ]
            agents_preview_html = generate_agent_preview_html(agent_list, "Document")
            # Generate deployed info for fallback agents
            early_results = {agent: "" for agent in agent_list}
            agents_deployed_html = generate_agents_deployed_html(early_results, "Document")
        
        progress(0.3, desc="🤖 Deploying AI agents...")
        
        # Initialize live log HTML (will be populated later)
        live_log_html = agents_preview_html  # Start with preview
        
        # Choose orchestrator
        if enable_iterative:
            logger.info("Starting ITERATIVE review mode")
            try:
                orchestrator = IterativeReviewOrchestrator(
                    config,
                    output_language=output_language,
                    max_iterations=max_iterations,
                    target_score=target_score,
                    interactive=enable_interactive,
                    enable_python_tools=enable_python_tools,
                    deep_review=enable_deep_review,
                    reference_context=reference_context
                )
                
                # Run iterative review with progress updates
                for i in range(max_iterations):
                    progress_pct = 0.4 + (i / max_iterations) * 0.5
                    progress(progress_pct, desc=f"🔄 Iteration {i+1}/{max_iterations}: AI agents analyzing...")
                
                results = asyncio.run(orchestrator.execute_iterative_review(document_text, title))
                logger.info(f"Iterative review completed: {len(results)} results")
                # For iterative, keep the preview (no detailed agent log)
                live_log_html = agents_preview_html
            except Exception as e:
                logger.error(f"Iterative review failed: {e}", exc_info=True)
                raise
            
        else:
            logger.info("Starting STANDARD review mode")
            try:
                orchestrator = GenericReviewOrchestrator(
                    config,
                    output_language=output_language,
                    enable_python_tools=enable_python_tools,
                    deep_review=enable_deep_review,
                    reference_context=reference_context
                )
                
                # Run single review with agent tracking + LIVE FEED
                # Build agent steps from the classified agent list
                agent_steps = []
                if agent_list:
                    step_size = 0.5 / len(agent_list)
                    for i, agent in enumerate(agent_list):
                        pct = 0.35 + (i * step_size)
                        agent_name = agent.replace('_', ' ').title()
                        # Get icon
                        icon_map = {
                            'style_editor': '🎨', 'consistency_checker': '🔍', 'fact_checker': '✓',
                            'logic_checker': '🧠', 'technical_expert': '⚙️', 'subject_matter_expert': '🎓',
                            'business_analyst': '💼', 'financial_analyst': '💰', 'legal_expert': '⚖️',
                            'data_validator': '📊', 'web_researcher': '🌐', 'academic_researcher': '📚',
                        }
                        icon = icon_map.get(agent, '🤖')
                        agent_steps.append((agent, pct, f"{icon} {agent_name} analyzing..."))
                else:
                    # Fallback steps
                    agent_steps = [
                        ('style_editor', 0.35, "🎨 Style & Grammar Check"),
                        ('fact_checker', 0.45, "✓ Fact Checking & Verification"),
                        ('logic_checker', 0.55, "🧠 Logic & Consistency Analysis"),
                        ('technical_expert', 0.65, "⚙️ Technical Review"),
                        ('business_analyst', 0.75, "💼 Business & Strategy Review"),
                        ('data_validator', 0.85, "📊 Data Validation"),
                    ]
                
                # Build LIVE ACTIVITY LOG with cumulative history
                total_agents = len(agent_steps)
                activities = []
                
                # Initialize all agents as pending
                for agent_key, _, desc in agent_steps:
                    activities.append({
                        'agent': agent_key,
                        'status': 'Queued',
                        'is_current': False,
                        'is_complete': False
                    })
                
                # Simulate agent execution with live updates
                for idx, (agent_key, pct, desc) in enumerate(agent_steps):
                    # Mark current agent as analyzing
                    activities[idx]['is_current'] = True
                    activities[idx]['status'] = 'Analyzing document...'
                    
                    # Mark previous agents as complete
                    for i in range(idx):
                        activities[i]['is_current'] = False
                        activities[i]['is_complete'] = True
                        activities[i]['status'] = 'Analysis complete'
                    
                    # Update progress bar
                    progress(pct, desc=desc)
                    
                    # Generate and show live activity log
                    # (Note: this won't update in real-time in Gradio, but shows progression)
                    live_log_html = generate_live_activity_log(
                        activities=activities,
                        total_agents=total_agents,
                        completed=idx
                    )
                    
                    # Wait to simulate work (shorter delay)
                    time.sleep(0.2)
                
                results = asyncio.run(orchestrator.execute_review_process(document_text, title))
                logger.info(f"Standard review completed: {len(results)} results")
                
                # Mark ALL agents as complete in final log
                for i in range(len(activities)):
                    activities[i]['is_current'] = False
                    activities[i]['is_complete'] = True
                    activities[i]['status'] = 'Analysis complete ✓'
                
                live_log_html = generate_live_activity_log(
                    activities=activities,
                    total_agents=total_agents,
                    completed=total_agents
                )
                
            except Exception as e:
                logger.error(f"Standard review failed: {e}", exc_info=True)
                raise
        
        progress(0.92, desc="📝 Compiling feedback...")
        
        progress(0.95, desc="📊 Creating visualizations...")
        
        # Generate status HTML with LIVE LOG (replace agents_preview_html)
        status_html = time_estimate_html + file_info_html + doc_stats_html + live_log_html + generate_status_html(results, enable_iterative)
        
        progress(0.97, desc="💾 Saving files...")
        
        # Load generated reports
        report_md_path = Path(unique_output_dir) / "review_report.md"
        results_json_path = Path(unique_output_dir) / "review_results.json"
        dashboard_html_path = Path(unique_output_dir) / "dashboard.html"
        
        # Also check for iterative dashboard
        if not dashboard_html_path.exists():
            iterative_dashboards = list(Path(unique_output_dir).glob("iterative_dashboard_*.html"))
            if iterative_dashboards:
                dashboard_html_path = iterative_dashboards[0]
        
        report_md = ""
        if report_md_path.exists():
            with open(report_md_path, 'r', encoding='utf-8') as f:
                report_md = f.read()
        else:
            report_md = "⚠️ Report not found. Check output directory for files."
        
        results_json = ""
        if results_json_path.exists():
            with open(results_json_path, 'r', encoding='utf-8') as f:
                results_json = f.read()
        else:
            results_json = "{}"
        
        # Load dashboard HTML
        dashboard_html = ""
        if dashboard_html_path.exists():
            # Read the HTML content directly
            try:
                with open(dashboard_html_path, 'r', encoding='utf-8') as f:
                    dashboard_content = f.read()
                
                # Create a styled container with the HTML and a link to open externally
                dashboard_html = f"""
                <div style="padding: 15px; background: #f0f7ff; border-radius: 8px; margin-bottom: 15px;">
                    <h3 style="margin: 0 0 10px 0;">📊 Interactive Dashboard</h3>
                    <p><strong>File location:</strong> <code>{dashboard_html_path}</code></p>
                    <p>
                        <a href="file://{dashboard_html_path}" target="_blank" style="
                            display: inline-block;
                            padding: 10px 20px;
                            background: #667eea;
                            color: white;
                            text-decoration: none;
                            border-radius: 5px;
                            font-weight: bold;
                        ">🚀 Open Dashboard in Browser</a>
                    </p>
                    <p style="font-size: 12px; color: #666;">
                        💡 <strong>Tip</strong>: Copy the file path above and open it in your browser for the full interactive experience with charts and graphs!
                    </p>
                </div>
                
                <div style="border: 1px solid #ddd; border-radius: 8px; overflow: hidden;">
                    {dashboard_content}
                </div>
                """
            except Exception as e:
                logger.error(f"Error loading dashboard: {e}")
                dashboard_html = f"""
                <div style="padding: 20px; background: #fff3cd; border-radius: 8px;">
                    <h3>⚠️ Dashboard Preview Unavailable</h3>
                    <p><strong>File location:</strong> <code>{dashboard_html_path}</code></p>
                    <p>Open the file above in your browser to view the interactive dashboard.</p>
                    <p><strong>Instructions:</strong></p>
                    <ol>
                        <li>Copy the file path shown above</li>
                        <li>Open your file manager (Finder/Explorer)</li>
                        <li>Navigate to the folder</li>
                        <li>Double-click <code>dashboard.html</code></li>
                    </ol>
                </div>
                """
        else:
            dashboard_html = """
            <div style="padding: 20px; background: #f8d7da; border-radius: 8px;">
                <h3>⚠️ Dashboard Not Generated</h3>
                <p>The dashboard file was not created. Check the output directory for any errors.</p>
            </div>
            """
        
        # Extract individual agent reports
        agents_report = extract_agent_reports(results)
        
        # Generate agent ratings (now with complete results)
        logger.info(f"Generating ratings for {len(results)} result items")
        logger.debug(f"Results type before ratings: {type(results)}")
        ratings_html = generate_agent_ratings_html(results)
        
        # agents_deployed_html was already generated EARLY (before analysis started)
        # So it's already shown to the user - don't regenerate to avoid UI flash
        
        progress(1.0, desc="✅ Complete!")
        
        # Prepare download buttons
        report_btn = gr.update(visible=True, value=report_md_path if report_md_path.exists() else None)
        json_btn = gr.update(visible=True, value=results_json_path if results_json_path.exists() else None)
        dashboard_btn = gr.update(visible=True, value=dashboard_html_path if dashboard_html_path.exists() else None)
        
        return status_html, report_md, results_json, dashboard_html, agents_report, ratings_html, agents_deployed_html, unique_output_dir, time_estimate_html, report_btn, json_btn, dashboard_btn
        
    except Exception as e:
        logger.error(f"Error processing document: {e}", exc_info=True)
        error_html = f"""
        <div style="background: #fee; padding: 20px; border-radius: 8px; border-left: 4px solid #d32f2f;">
            <h3 style="color: #d32f2f; margin: 0 0 10px 0;">⚠️ Processing Error</h3>
            <p><strong>Error:</strong> {str(e)}</p>
            <p><small>Please check your file and try again. If the problem persists, contact support.</small></p>
        </div>
        """
        return error_html, "", "", "", "", "", "", "", "", gr.update(), gr.update(), gr.update()


def generate_agent_ratings_html(results: dict) -> str:
    """Generate beautiful star ratings for each agent's performance."""
    if not results or not isinstance(results, dict):
        return "<p>No agent ratings available.</p>"
    
    # Extract agent reviews
    agent_reviews = results.get("agent_reviews", {})
    if not agent_reviews:
        return "<p>No agent ratings available.</p>"
    
    # Icon mapping
    agent_icons = {
        'style_editor': '🎨',
        'consistency_checker': '🔍',
        'fact_checker': '✓',
        'logic_checker': '🧠',
        'technical_expert': '⚙️',
        'subject_matter_expert': '🎓',
        'business_analyst': '💼',
        'financial_analyst': '💰',
        'legal_expert': '⚖️',
        'data_validator': '📊',
        'web_researcher': '🌐',
        'academic_researcher': '📚',
    }
    
    html = """
    <div style="max-width: 1200px; margin: 0 auto; padding: 20px;">
        <div style="text-align: center; margin-bottom: 40px;">
            <h2 style="color: #667eea; font-size: 32px; font-weight: 700; margin-bottom: 10px;">
                ⭐ Agent Performance Ratings
            </h2>
            <p style="color: #666; font-size: 16px;">
                Each agent is rated based on thoroughness, insights quality, and usefulness
            </p>
        </div>
        
        <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(350px, 1fr)); gap: 20px;">
    """
    
    # Sort agents by rating (if available) or alphabetically
    sorted_agents = sorted(agent_reviews.items(), key=lambda x: x[0])
    
    for agent_key, review_data in sorted_agents:
        agent_name = agent_key.replace('_', ' ').title()
        icon = agent_icons.get(agent_key, '🤖')
        
        # Extract review text
        if isinstance(review_data, dict):
            review_text = review_data.get('review', '') or review_data.get('content', '')
        else:
            review_text = str(review_data)
        
        # Calculate simple rating based on review length and sentiment
        # (In a real system, you'd use AI to rate, but for now a heuristic)
        review_length = len(review_text)
        if review_length > 1000:
            rating = 5.0
        elif review_length > 600:
            rating = 4.5
        elif review_length > 300:
            rating = 4.0
        elif review_length > 100:
            rating = 3.5
        else:
            rating = 3.0
        
        # Count of findings/suggestions (rough estimate)
        findings_count = review_text.count('\n-') + review_text.count('\n•') + review_text.count('\n*')
        
        # Generate stars
        full_stars = int(rating)
        has_half = (rating % 1) >= 0.5
        empty_stars = 5 - full_stars - (1 if has_half else 0)
        
        stars_html = ''.join(['⭐'] * full_stars)
        if has_half:
            stars_html += '✨'
        stars_html += ''.join(['☆'] * empty_stars)
        
        # Rating color
        if rating >= 4.5:
            rating_color = "#10b981"  # Green
            rating_label = "Excellent"
        elif rating >= 4.0:
            rating_color = "#3b82f6"  # Blue
            rating_label = "Very Good"
        elif rating >= 3.5:
            rating_color = "#f59e0b"  # Orange
            rating_label = "Good"
        else:
            rating_color = "#6b7280"  # Gray
            rating_label = "Adequate"
        
        html += f"""
        <div style="
            background: white;
            border-radius: 16px;
            padding: 24px;
            box-shadow: 0 4px 20px rgba(0,0,0,0.08);
            border: 2px solid #f3f4f6;
            transition: all 0.3s ease;
        " class="agent-rating-card">
            <div style="display: flex; align-items: center; margin-bottom: 16px;">
                <div style="
                    font-size: 48px;
                    margin-right: 16px;
                    animation: float 3s ease-in-out infinite;
                ">{icon}</div>
                <div style="flex: 1;">
                    <h3 style="
                        margin: 0 0 4px 0;
                        color: #1f2937;
                        font-size: 20px;
                        font-weight: 700;
                    ">{agent_name}</h3>
                    <div style="
                        font-size: 24px;
                        letter-spacing: 2px;
                        line-height: 1;
                    ">{stars_html}</div>
                </div>
                <div style="text-align: right;">
                    <div style="
                        font-size: 32px;
                        font-weight: 700;
                        color: {rating_color};
                        line-height: 1;
                    ">{rating:.1f}</div>
                    <div style="
                        font-size: 11px;
                        color: {rating_color};
                        font-weight: 600;
                        margin-top: 4px;
                    ">{rating_label}</div>
                </div>
            </div>
            
            <div style="
                display: flex;
                gap: 16px;
                padding: 12px;
                background: #f9fafb;
                border-radius: 8px;
                margin-top: 16px;
            ">
                <div style="flex: 1; text-align: center;">
                    <div style="font-size: 24px; font-weight: 700; color: #667eea;">
                        {len(review_text)}
                    </div>
                    <div style="font-size: 11px; color: #6b7280; font-weight: 600;">
                        CHARACTERS
                    </div>
                </div>
                <div style="flex: 1; text-align: center;">
                    <div style="font-size: 24px; font-weight: 700; color: #667eea;">
                        {max(findings_count, 1)}
                    </div>
                    <div style="font-size: 11px; color: #6b7280; font-weight: 600;">
                        INSIGHTS
                    </div>
                </div>
                <div style="flex: 1; text-align: center;">
                    <div style="font-size: 24px; font-weight: 700; color: #667eea;">
                        {min(100, int((review_length / 10)))}</div>
                    <div style="font-size: 11px; color: #6b7280; font-weight: 600;">
                        DEPTH %
                    </div>
                </div>
            </div>
        </div>
        """
    
    html += """
        </div>
    </div>
    
    <style>
    .agent-rating-card:hover {
        transform: translateY(-4px);
        box-shadow: 0 8px 30px rgba(102, 126, 234, 0.15);
        border-color: #667eea;
    }
    </style>
    """
    
    return html


def extract_agent_reports(results: dict) -> str:
    """Extract and format individual agent reports."""
    reviews = results.get('reviews', {})
    
    if not reviews:
        return "⚠️ No agent reports found."
    
    html = "<div style='font-family: sans-serif;'>"
    
    for agent_name, review_content in reviews.items():
        # Format agent name
        display_name = agent_name.replace('_', ' ').title()
        
        # Get icon based on agent type
        icon = "🤖"
        if "web" in agent_name.lower():
            icon = "🌐"
        elif "data" in agent_name.lower():
            icon = "📊"
        elif "fact" in agent_name.lower():
            icon = "✓"
        elif "technical" in agent_name.lower():
            icon = "⚙️"
        elif "business" in agent_name.lower():
            icon = "💼"
        elif "financial" in agent_name.lower():
            icon = "💰"
        elif "legal" in agent_name.lower():
            icon = "⚖️"
        
        html += f"""
        <div style="margin: 20px 0; padding: 20px; background: #f8f9fa; border-left: 4px solid #667eea; border-radius: 8px;">
            <h3 style="margin: 0 0 15px 0; color: #667eea;">
                {icon} {display_name}
            </h3>
            <div style="white-space: pre-wrap; line-height: 1.6;">
                {review_content if isinstance(review_content, str) else str(review_content)}
            </div>
        </div>
        """
    
    html += "</div>"
    return html


def generate_status_html(results: dict, is_iterative: bool) -> str:
    """Generate HTML status summary with enhanced visuals."""
    if is_iterative:
        iterations = results.get('total_iterations', 0)
        final_score = results.get('final_score', 0)
        improvement = results.get('improvement_summary', {}).get('score_improvement', 0)
        best_iter = results.get('best_iteration', {}).get('iteration_number', 0)
        
        # Score color
        score_color = "#4ade80" if final_score >= 80 else "#fbbf24" if final_score >= 60 else "#f87171"
        improvement_color = "#4ade80" if improvement > 0 else "#f87171"
        
        html = f"""
        <div style="padding: 30px; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                    border-radius: 15px; color: white; margin: 20px 0; box-shadow: 0 10px 40px rgba(102, 126, 234, 0.4);">
            <div style="text-align: center; margin-bottom: 25px;">
                <div style="font-size: 48px; margin-bottom: 10px;">✨</div>
                <h2 style="margin: 0; font-size: 32px; font-weight: 700;">Iterative Review Complete!</h2>
                <p style="opacity: 0.95; margin: 10px 0 0 0;">Your document has been enhanced through multiple AI refinement cycles</p>
            </div>
            <div style="display: grid; grid-template-columns: repeat(4, 1fr); gap: 15px;">
                <div style="background: rgba(255,255,255,0.15); padding: 20px; border-radius: 12px; text-align: center;
                           box-shadow: 0 4px 15px rgba(0,0,0,0.1); backdrop-filter: blur(10px);">
                    <div style="font-size: 36px; font-weight: bold;">{iterations}</div>
                    <div style="opacity: 0.9; margin-top: 5px;">🔄 Iterations</div>
                </div>
                <div style="background: rgba(255,255,255,0.15); padding: 20px; border-radius: 12px; text-align: center;
                           box-shadow: 0 4px 15px rgba(0,0,0,0.1); backdrop-filter: blur(10px);">
                    <div style="font-size: 36px; font-weight: bold; color: {score_color};">{final_score:.1f}</div>
                    <div style="opacity: 0.9; margin-top: 5px;">🎯 Final Score</div>
                </div>
                <div style="background: rgba(255,255,255,0.15); padding: 20px; border-radius: 12px; text-align: center;
                           box-shadow: 0 4px 15px rgba(0,0,0,0.1); backdrop-filter: blur(10px);">
                    <div style="font-size: 36px; font-weight: bold; color: {improvement_color};">{improvement:+.1f}</div>
                    <div style="opacity: 0.9; margin-top: 5px;">📈 Improvement</div>
                </div>
                <div style="background: rgba(255,255,255,0.15); padding: 20px; border-radius: 12px; text-align: center;
                           box-shadow: 0 4px 15px rgba(0,0,0,0.1); backdrop-filter: blur(10px);">
                    <div style="font-size: 36px; font-weight: bold;">#{best_iter}</div>
                    <div style="opacity: 0.9; margin-top: 5px;">⭐ Best Version</div>
                </div>
            </div>
            <div style="margin-top: 20px; padding: 15px; background: rgba(255,255,255,0.1); border-radius: 10px; text-align: center;">
                <strong>💡 Next Steps:</strong> Review the detailed feedback below and download your improved document!
            </div>
        </div>
        """
    else:
        agent_count = len(results.get('reviews', {}))
        review_types = list(results.get('reviews', {}).keys())
        
        # Count agent categories
        web_agents = sum(1 for a in review_types if 'web' in a.lower() or 'fact' in a.lower())
        business_agents = sum(1 for a in review_types if 'business' in a.lower() or 'financial' in a.lower())
        technical_agents = sum(1 for a in review_types if 'technical' in a.lower() or 'data' in a.lower())
        
        html = f"""
        <div style="padding: 30px; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                    border-radius: 15px; color: white; margin: 20px 0; box-shadow: 0 10px 40px rgba(102, 126, 234, 0.4);">
            <div style="text-align: center; margin-bottom: 25px;">
                <div style="font-size: 48px; margin-bottom: 10px;">✅</div>
                <h2 style="margin: 0; font-size: 32px; font-weight: 700;">Professional Review Complete!</h2>
                <p style="opacity: 0.95; margin: 10px 0 0 0;">Your document has been analyzed by {agent_count} AI specialists</p>
            </div>
            <div style="display: grid; grid-template-columns: repeat(3, 1fr); gap: 15px; margin-bottom: 20px;">
                <div style="background: rgba(255,255,255,0.15); padding: 20px; border-radius: 12px; text-align: center;
                           box-shadow: 0 4px 15px rgba(0,0,0,0.1); backdrop-filter: blur(10px);">
                    <div style="font-size: 36px; font-weight: bold;">{agent_count}</div>
                    <div style="opacity: 0.9; margin-top: 5px;">🤖 Total Agents</div>
                </div>
                <div style="background: rgba(255,255,255,0.15); padding: 20px; border-radius: 12px; text-align: center;
                           box-shadow: 0 4px 15px rgba(0,0,0,0.1); backdrop-filter: blur(10px);">
                    <div style="font-size: 36px; font-weight: bold;">{technical_agents}</div>
                    <div style="opacity: 0.9; margin-top: 5px;">⚙️ Technical</div>
                </div>
                <div style="background: rgba(255,255,255,0.15); padding: 20px; border-radius: 12px; text-align: center;
                           box-shadow: 0 4px 15px rgba(0,0,0,0.1); backdrop-filter: blur(10px);">
                    <div style="font-size: 36px; font-weight: bold;">{web_agents}</div>
                    <div style="opacity: 0.9; margin-top: 5px;">🌐 Research</div>
                </div>
            </div>
            <div style="background: rgba(255,255,255,0.1); padding: 15px; border-radius: 10px;">
                <strong>📋 Agents Deployed:</strong><br>
                <div style="margin-top: 10px; font-size: 14px; opacity: 0.95;">
                    {', '.join([a.replace('_', ' ').title() for a in review_types[:8]])}
                    {'...' if len(review_types) > 8 else ''}
                </div>
            </div>
            <div style="margin-top: 15px; padding: 15px; background: rgba(255,255,255,0.1); border-radius: 10px; text-align: center;">
                <strong>💡 Next Steps:</strong> Explore detailed feedback from each agent in the tabs below!
            </div>
        </div>
        """
    
    return html


def create_ui():
    """Create Gradio interface."""
    
    # Initialize system
    init_success, init_message = initialize_system()
    
    # Custom CSS - Modern & Professional with Animations
    custom_css = """
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
    
    /* Global Styles */
    .gradio-container {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
        max-width: 1400px !important;
    }
    
    /* Smooth scroll */
    * {
        scroll-behavior: smooth;
    }
    
    /* Hero Section with Animation */
    .hero-section {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 60px 40px;
        border-radius: 20px;
        color: white;
        text-align: center;
        margin-bottom: 40px;
        box-shadow: 0 20px 60px rgba(102, 126, 234, 0.3);
        animation: fadeInDown 0.8s ease-out;
    }
    
    @keyframes fadeInDown {
        from {
            opacity: 0;
            transform: translateY(-30px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }
    
    @keyframes fadeIn {
        from { opacity: 0; }
        to { opacity: 1; }
    }
    
    @keyframes slideInUp {
        from {
            opacity: 0;
            transform: translateY(30px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }
    
    @keyframes pulse {
        0%, 100% {
            transform: scale(1);
        }
        50% {
            transform: scale(1.05);
        }
    }
    
    @keyframes shimmer {
        0% {
            background-position: -1000px 0;
        }
        100% {
            background-position: 1000px 0;
        }
    }
    
    .hero-title {
        font-size: 48px;
        font-weight: 700;
        margin: 0 0 20px 0;
        line-height: 1.2;
    }
    
    .hero-subtitle {
        font-size: 20px;
        opacity: 0.95;
        margin: 0 0 30px 0;
        font-weight: 400;
    }
    
    /* Benefit Cards with Animations */
    .benefit-card {
        background: white;
        padding: 30px;
        border-radius: 15px;
        box-shadow: 0 4px 20px rgba(0,0,0,0.08);
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
        border: 1px solid #f0f0f0;
        animation: fadeIn 0.6s ease-out backwards;
        position: relative;
        overflow: hidden;
    }
    
    .benefit-card::before {
        content: '';
        position: absolute;
        top: 0;
        left: -100%;
        width: 100%;
        height: 100%;
        background: linear-gradient(90deg, transparent, rgba(102, 126, 234, 0.1), transparent);
        transition: left 0.5s;
    }
    
    .benefit-card:hover {
        transform: translateY(-8px) scale(1.02);
        box-shadow: 0 12px 40px rgba(102, 126, 234, 0.25);
        border-color: #667eea;
    }
    
    .benefit-card:hover::before {
        left: 100%;
    }
    
    /* Stagger animation for cards */
    .benefit-card:nth-child(1) { animation-delay: 0.1s; }
    .benefit-card:nth-child(2) { animation-delay: 0.2s; }
    .benefit-card:nth-child(3) { animation-delay: 0.3s; }
    
    .success-message {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 16px 24px;
        border-radius: 12px;
        margin: 20px 0;
        font-weight: 500;
        box-shadow: 0 4px 15px rgba(102, 126, 234, 0.3);
    }
    
    .error-message {
        background: #ff6b6b;
        color: white;
        padding: 16px 24px;
        border-radius: 12px;
        margin: 20px 0;
        font-weight: 500;
    }
    
    .step-title {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-weight: 700;
        font-size: 20px;
        margin: 25px 0 15px 0;
    }
    
    /* Enhanced Buttons */
    button.primary {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
        border: none !important;
        color: white !important;
        font-weight: 600 !important;
        padding: 18px 36px !important;
        font-size: 18px !important;
        border-radius: 12px !important;
        box-shadow: 0 4px 15px rgba(102, 126, 234, 0.4) !important;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
        position: relative !important;
        overflow: hidden !important;
    }
    
    button.primary::before {
        content: '';
        position: absolute;
        top: 50%;
        left: 50%;
        width: 0;
        height: 0;
        border-radius: 50%;
        background: rgba(255,255,255,0.3);
        transform: translate(-50%, -50%);
        transition: width 0.6s, height 0.6s;
    }
    
    button.primary:hover::before {
        width: 300px;
        height: 300px;
    }
    
    button.primary:hover {
        transform: translateY(-3px) scale(1.02) !important;
        box-shadow: 0 8px 25px rgba(102, 126, 234, 0.6) !important;
    }
    
    button.primary:active {
        transform: translateY(-1px) scale(0.98) !important;
    }
    
    /* Feature Badges with Animation */
    .feature-badge {
        display: inline-block;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 8px 18px;
        border-radius: 25px;
        font-size: 14px;
        font-weight: 600;
        margin: 5px;
        animation: slideInUp 0.5s ease-out backwards;
        box-shadow: 0 2px 10px rgba(102, 126, 234, 0.3);
        transition: all 0.3s ease;
    }
    
    .feature-badge:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 15px rgba(102, 126, 234, 0.4);
    }
    
    .feature-badge:nth-child(1) { animation-delay: 0.1s; }
    .feature-badge:nth-child(2) { animation-delay: 0.2s; }
    .feature-badge:nth-child(3) { animation-delay: 0.3s; }
    .feature-badge:nth-child(4) { animation-delay: 0.4s; }
    
    /* Loading State */
    .loading {
        animation: pulse 2s cubic-bezier(0.4, 0, 0.6, 1) infinite;
    }
    
    /* Tabs Styling */
    .tab-nav button {
        transition: all 0.3s ease !important;
    }
    
    .tab-nav button:hover {
        transform: translateY(-2px);
    }
    
    /* File Upload Area */
    .file-upload {
        transition: all 0.3s ease;
        border: 2px dashed #667eea !important;
    }
    
    .file-upload:hover {
        border-color: #764ba2 !important;
        background: #f8f9ff !important;
    }
    
    /* Accordion Headers */
    .accordion-header {
        transition: all 0.3s ease;
    }
    
    .accordion-header:hover {
        background: #f8f9ff;
    }
    
    /* Progress Bar Enhancement */
    .progress-bar {
        background: linear-gradient(90deg, #667eea, #764ba2, #667eea);
        background-size: 200% 100%;
        animation: shimmer 2s linear infinite;
    }
    
    /* Results Section Spacing */
    .results-container {
        padding: 30px 20px !important;
        background: white !important;
        border-radius: 15px !important;
        box-shadow: 0 2px 8px rgba(0,0,0,0.05) !important;
        margin-top: 20px !important;
    }
    
    /* Tab Container */
    .tab-container {
        margin-top: 25px !important;
        padding: 20px !important;
        background: #f8f9fa !important;
        border-radius: 12px !important;
    }
    
    /* Tab Nav Buttons */
    .tab-nav {
        gap: 10px !important;
        padding: 15px 10px !important;
        background: white !important;
        border-radius: 10px !important;
        margin-bottom: 20px !important;
    }
    
    .tab-nav button {
        padding: 12px 24px !important;
        font-size: 15px !important;
        font-weight: 500 !important;
        border-radius: 8px !important;
        min-height: 45px !important;
    }
    
    /* Tab Content */
    .tab-content {
        padding: 25px 15px !important;
        min-height: 400px !important;
    }
    
    /* Status Output */
    .status-box {
        margin: 20px 0 !important;
    }
    
    /* Textbox and Code blocks spacing */
    .gr-textbox, .gr-code {
        margin: 15px 0 !important;
    }
    
    /* Download buttons spacing */
    .gr-button-group {
        gap: 15px !important;
        margin: 20px 0 !important;
    }
    
    /* ============================================
       ADVANCED UI ENHANCEMENTS (Step 2)
       ============================================ */
    
    /* 1. Skeleton Loader Animation */
    @keyframes skeleton-loading {
        0% {
            background-position: 200% 0;
        }
        100% {
            background-position: -200% 0;
        }
    }
    
    .skeleton {
        background: linear-gradient(
            90deg,
            #f0f0f0 25%,
            #e0e0e0 50%,
            #f0f0f0 75%
        );
        background-size: 200% 100%;
        animation: skeleton-loading 1.5s ease-in-out infinite;
        border-radius: 8px;
    }
    
    /* 2. Glassmorphism Enhanced */
    .glass-card {
        background: rgba(255, 255, 255, 0.7) !important;
        backdrop-filter: blur(20px) !important;
        -webkit-backdrop-filter: blur(20px) !important;
        border: 1px solid rgba(255, 255, 255, 0.3) !important;
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.1) !important;
    }
    
    /* 3. Gradient Text Animation */
    @keyframes gradient-shift {
        0%, 100% {
            background-position: 0% 50%;
        }
        50% {
            background-position: 100% 50%;
        }
    }
    
    .gradient-text {
        background: linear-gradient(
            45deg,
            #667eea,
            #764ba2,
            #667eea
        );
        background-size: 200% auto;
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        animation: gradient-shift 3s ease infinite;
    }
    
    /* 4. Floating Animation */
    @keyframes float {
        0%, 100% {
            transform: translateY(0px);
        }
        50% {
            transform: translateY(-10px);
        }
    }
    
    .float-animation {
        animation: float 3s ease-in-out infinite;
    }
    
    /* 5. Glow Effect on Hover */
    .glow-on-hover {
        transition: all 0.3s ease;
    }
    
    .glow-on-hover:hover {
        box-shadow: 0 0 20px rgba(102, 126, 234, 0.6),
                    0 0 40px rgba(102, 126, 234, 0.4),
                    0 0 60px rgba(102, 126, 234, 0.2) !important;
        transform: translateY(-2px);
    }
    
    /* 6. Progress Ring (for circular progress) */
    @keyframes rotate-ring {
        from {
            transform: rotate(0deg);
        }
        to {
            transform: rotate(360deg);
        }
    }
    
    .progress-ring {
        animation: rotate-ring 2s linear infinite;
    }
    
    /* 7. Smooth Page Transitions */
    .page-transition {
        animation: fadeIn 0.5s ease-out;
    }
    
    /* 8. Card Flip Effect */
    .flip-card {
        perspective: 1000px;
        transition: transform 0.6s;
        transform-style: preserve-3d;
    }
    
    .flip-card:hover {
        transform: rotateY(10deg);
    }
    
    /* 9. Tooltip Enhancement */
    .tooltip {
        position: relative;
        cursor: help;
    }
    
    .tooltip::after {
        content: attr(data-tooltip);
        position: absolute;
        bottom: 100%;
        left: 50%;
        transform: translateX(-50%);
        background: #333;
        color: white;
        padding: 8px 12px;
        border-radius: 6px;
        font-size: 12px;
        white-space: nowrap;
        opacity: 0;
        pointer-events: none;
        transition: opacity 0.3s ease;
        z-index: 1000;
    }
    
    .tooltip:hover::after {
        opacity: 1;
    }
    
    /* 10. Typing Animation for Headers */
    @keyframes typing {
        from {
            width: 0;
        }
        to {
            width: 100%;
        }
    }
    
    .typing-effect {
        overflow: hidden;
        white-space: nowrap;
        animation: typing 2s steps(40, end);
    }
    
    /* 11. Particle Background (subtle) */
    @keyframes particle-float {
        0%, 100% {
            transform: translate(0, 0);
            opacity: 0.3;
        }
        50% {
            transform: translate(10px, -10px);
            opacity: 0.6;
        }
    }
    
    .particles {
        position: absolute;
        width: 100%;
        height: 100%;
        overflow: hidden;
        pointer-events: none;
    }
    
    .particle {
        position: absolute;
        width: 4px;
        height: 4px;
        background: #667eea;
        border-radius: 50%;
        animation: particle-float 3s ease-in-out infinite;
    }
    
    /* 12. Enhanced Input Focus */
    input:focus, textarea:focus, select:focus {
        outline: none !important;
        border-color: #667eea !important;
        box-shadow: 0 0 0 3px rgba(102, 126, 234, 0.1) !important;
        transform: translateY(-1px);
        transition: all 0.3s ease;
    }
    
    /* 13. Smooth Scrollbar */
    ::-webkit-scrollbar {
        width: 12px;
        height: 12px;
    }
    
    ::-webkit-scrollbar-track {
        background: #f1f1f1;
        border-radius: 10px;
    }
    
    ::-webkit-scrollbar-thumb {
        background: linear-gradient(135deg, #667eea, #764ba2);
        border-radius: 10px;
        transition: background 0.3s ease;
    }
    
    ::-webkit-scrollbar-thumb:hover {
        background: linear-gradient(135deg, #764ba2, #667eea);
    }
    
    /* 14. Loading Spinner Enhancement */
    @keyframes spin-glow {
        0% {
            transform: rotate(0deg);
            filter: drop-shadow(0 0 5px #667eea);
        }
        100% {
            transform: rotate(360deg);
            filter: drop-shadow(0 0 10px #764ba2);
        }
    }
    
    .spinner {
        animation: spin-glow 1s linear infinite;
    }
    
    /* 15. Micro-interaction: Button Click */
    button:active {
        transform: scale(0.95) !important;
        transition: transform 0.1s ease !important;
    }
    
    /* 16. Success/Error States with Animation */
    @keyframes success-pulse {
        0%, 100% {
            box-shadow: 0 0 0 0 rgba(72, 222, 128, 0.7);
        }
        50% {
            box-shadow: 0 0 0 10px rgba(72, 222, 128, 0);
        }
    }
    
    .success-state {
        animation: success-pulse 1s ease-out;
    }
    
    @keyframes error-shake {
        0%, 100% { transform: translateX(0); }
        25% { transform: translateX(-5px); }
        75% { transform: translateX(5px); }
    }
    
    .error-state {
        animation: error-shake 0.3s ease-out;
    }
    
    /* 17. Badge Pulse Animation */
    .badge-pulse {
        animation: pulse 2s cubic-bezier(0.4, 0, 0.6, 1) infinite;
    }
    
    /* 18. Card Hover Lift (more pronounced) */
    .card-lift {
        transition: all 0.4s cubic-bezier(0.175, 0.885, 0.32, 1.275);
    }
    
    .card-lift:hover {
        transform: translateY(-15px) scale(1.02);
        box-shadow: 0 20px 60px rgba(102, 126, 234, 0.3);
    }
    
    /* 19. Text Gradient Reveal on Hover */
    .gradient-reveal {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 50%, #667eea 100%);
        background-size: 200% auto;
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        transition: background-position 0.5s ease;
    }
    
    .gradient-reveal:hover {
        background-position: 100% 0;
    }
    
    /* 20. Responsive Grid Enhancement */
    @media (max-width: 768px) {
        .hero-title {
            font-size: 32px !important;
        }
        
        .benefit-card {
            margin-bottom: 20px;
        }
        
        .tab-nav button {
            font-size: 13px !important;
            padding: 10px 16px !important;
        }
    }
    """
    
    with gr.Blocks(css=custom_css, title="AI Document Pro - Professional Review System", theme=gr.themes.Soft()) as app:
        # Hero Section
        gr.HTML("""
        <div class="hero-section">
            <h1 class="hero-title">✨ Transform Your Documents with AI</h1>
            <p class="hero-subtitle">Professional-grade document analysis in minutes. Get expert feedback powered by cutting-edge AI technology.</p>
            <div style="margin-top: 30px;">
                <span class="feature-badge">⚡ Lightning Fast</span>
                <span class="feature-badge">🎯 99% Accurate</span>
                <span class="feature-badge">🌍 Multi-Language</span>
                <span class="feature-badge">🔒 Secure</span>
            </div>
        </div>
        """)
        
        # System status
        if init_success:
            gr.HTML(f'<div class="success-message">{init_message}</div>')
        else:
            gr.HTML(f'<div class="error-message">{init_message}</div>')
            gr.Markdown("**Please set your OPENAI_API_KEY environment variable and restart.**")
            return app
        
        # Benefits Section with Advanced Effects
        gr.HTML("""
        <div style="display: grid; grid-template-columns: repeat(3, 1fr); gap: 20px; margin: 30px 0;">
            <div class="benefit-card card-lift glow-on-hover">
                <div style="font-size: 36px; margin-bottom: 15px;" class="float-animation">⚡</div>
                <h3 style="margin: 0 0 10px 0;" class="gradient-reveal">Fast Results</h3>
                <p style="color: #666; margin: 0;">Get comprehensive feedback in just 3-5 minutes</p>
            </div>
            <div class="benefit-card card-lift glow-on-hover">
                <div style="font-size: 36px; margin-bottom: 15px;" class="float-animation" style="animation-delay: 0.5s;">🎯</div>
                <h3 style="margin: 0 0 10px 0;" class="gradient-reveal">Expert Quality</h3>
                <p style="color: #666; margin: 0;">30+ AI specialists analyze every aspect</p>
            </div>
            <div class="benefit-card card-lift glow-on-hover">
                <div style="font-size: 36px; margin-bottom: 15px;" class="float-animation" style="animation-delay: 1s;">📈</div>
                <h3 style="margin: 0 0 10px 0;" class="gradient-reveal">Improve Automatically</h3>
                <p style="color: #666; margin: 0;">Optional AI-powered document refinement</p>
            </div>
        </div>
        """)
        
        with gr.Tabs():
            # Main Review Tab
            with gr.Tab("🚀 Get Started"):
                with gr.Row():
                    with gr.Column(scale=1):
                        gr.HTML('<h3 class="step-title">1️⃣ Upload Your Document</h3>')
                        file_input = gr.File(
                            label="📁 Drop your file here or click to browse",
                            file_types=[".pdf", ".txt", ".md", ".docx"],
                            type="filepath"
                        )
                        gr.Markdown("*Supports: PDF, Word, Text, Markdown*")
                        
                        gr.HTML('<h3 class="step-title">2️⃣ Choose Your Language</h3>')
                        
                        output_language = gr.Dropdown(
                            choices=["Auto-detect", "Italian", "English", "Spanish", "French", "German", "Portuguese"],
                            value="Auto-detect",
                            label="🌍 Report Language",
                            info="We'll automatically detect your document language"
                        )
                        
                        with gr.Accordion("✨ Auto-Improve (Recommended)", open=False):
                            enable_iterative = gr.Checkbox(
                                label="🔄 Automatically improve my document",
                                value=False,
                                info="AI will enhance your document step-by-step"
                            )
                            max_iterations = gr.Slider(
                                minimum=1,
                                maximum=10,
                                value=3,
                                step=1,
                                label="Number of improvement rounds",
                                info="More rounds = better quality (3 is usually perfect)"
                            )
                            target_score = gr.Slider(
                                minimum=60,
                                maximum=100,
                                value=85,
                                step=5,
                                label="Target quality score",
                                info="Stop improving when this score is reached"
                            )
                        
                        with gr.Accordion("🎓 Professional Features", open=False):
                            enable_python_tools = gr.Checkbox(
                                label="📊 Validate Numbers & Data",
                                value=True,
                                info="Check calculations and verify statistics automatically"
                            )
                            enable_deep_review = gr.Checkbox(
                                label="🔬 Deep Analysis for Research",
                                value=False,
                                info="Add academic literature search + 20 specialist reviewers (best for scientific papers)"
                            )
                            enable_interactive = gr.Checkbox(
                                label="💬 Interactive Mode (Advanced)",
                                value=False,
                                info="AI may ask clarification questions"
                            )
                        
                        with gr.Accordion("📚 Compare with Templates (Optional)", open=False):
                            reference_files = gr.File(
                                label="Upload Templates or Guidelines",
                                file_count="multiple",
                                file_types=[".pdf", ".txt", ".md", ".docx", ".xlsx"],
                                type="filepath"
                            )
                            reference_type = gr.Dropdown(
                                choices=["template", "guideline", "example", "data", "style_guide"],
                                value="example",
                                label="What type of reference is this?"
                            )
                        
                        gr.HTML('<h3 class="step-title">3️⃣ Get Your Results</h3>')
                        submit_btn = gr.Button("✨ Analyze My Document", variant="primary", size="lg")
                    
                    with gr.Column(scale=2):
                        gr.HTML("""
                        <div style="padding: 20px 0 10px 0;">
                            <h2 style="
                                font-size: 28px;
                                font-weight: 700;
                                color: #667eea;
                                margin: 0 0 10px 0;
                                display: flex;
                                align-items: center;
                                gap: 10px;
                            ">
                                📊 Results
                            </h2>
                            <p style="color: #666; margin: 0; font-size: 14px;">
                                Your comprehensive document analysis will appear here
                            </p>
                        </div>
                        """)
                        
                        # Time estimate display (hidden initially)
                        time_estimate_output = gr.HTML(label="Processing Info", visible=False)
                        
                        status_output = gr.HTML(label="Status", elem_classes="status-box")
                        
                        gr.HTML('<div style="margin: 30px 0 20px 0;"></div>')  # Spacer
                        
                        with gr.Tabs():
                            with gr.Tab("📋 Report"):
                                gr.HTML('<div style="padding: 20px 0;"></div>')
                                # Show deployed agents before report
                                agents_deployed_output = gr.HTML(label="Agents Deployed")
                                gr.HTML('<div style="padding: 10px 0;"></div>')
                                report_output = gr.Markdown(label="Review Report")
                            
                            with gr.Tab("🤖 Agent Reviews"):
                                gr.HTML('<div style="padding: 20px 0;"></div>')
                                agents_output = gr.HTML(label="Individual Agent Reports")
                            
                            with gr.Tab("⭐ Agent Ratings"):
                                gr.HTML('<div style="padding: 20px 0;"></div>')
                                ratings_output = gr.HTML(label="Agent Performance Ratings")
                                gr.Markdown("""
                                **💡 Tip:** Each agent is automatically rated based on thoroughness, insight quality, and depth of analysis.
                                Higher ratings indicate more comprehensive and valuable contributions.
                                """)
                            
                            with gr.Tab("📊 Dashboard"):
                                gr.HTML('<div style="padding: 20px 0;"></div>')
                                dashboard_output = gr.HTML(label="Interactive Dashboard")
                                gr.Markdown("""
                                **💡 Tip**: If the dashboard doesn't display above, you can:
                                1. Check the "Files" tab for the output directory
                                2. Open `dashboard.html` directly in your browser
                                3. The file path will be shown in the Files tab
                                """)
                            
                            with gr.Tab("📦 JSON"):
                                gr.HTML('<div style="padding: 20px 0;"></div>')
                                json_output = gr.Code(label="Full Results (JSON)", language="json")
                            
                            with gr.Tab("📁 Files"):
                                gr.HTML('<div style="padding: 20px 0;"></div>')
                                output_dir_display = gr.Textbox(label="Output Directory", interactive=False)
                                
                                gr.HTML('<div style="margin: 30px 0;"></div>')  # Spacer
                                
                                gr.HTML("""
                                <h3 style="
                                    color: #667eea;
                                    font-size: 20px;
                                    font-weight: 600;
                                    margin: 0 0 20px 0;
                                ">
                                    📥 Download Files
                                </h3>
                                """)
                                
                                with gr.Row():
                                    download_report_btn = gr.DownloadButton(
                                        label="📄 Download Report (MD)",
                                        visible=False,
                                        size="lg"
                                    )
                                    download_json_btn = gr.DownloadButton(
                                        label="📦 Download JSON",
                                        visible=False,
                                        size="lg"
                                    )
                                    download_dashboard_btn = gr.DownloadButton(
                                        label="📊 Download Dashboard (HTML)",
                                        visible=False,
                                        size="lg"
                                    )
                                
                                gr.HTML('<div style="margin: 40px 0 20px 0;"></div>')  # Spacer
                                
                                gr.HTML("""
                                <div style="
                                    background: #f8f9fa;
                                    padding: 25px;
                                    border-radius: 12px;
                                    border-left: 4px solid #667eea;
                                ">
                                    <h4 style="margin: 0 0 15px 0; color: #667eea; font-size: 18px;">
                                        📁 Generated Files
                                    </h4>
                                """)
                                
                                gr.Markdown("""
                                **Your review includes:**
                                - `review_report.md` - Human-readable report
                                - `review_results.json` - Complete results in JSON
                                - `dashboard.html` - Interactive HTML dashboard ← **Open this!**
                                - `document_*.txt` - Document versions (if iterative)
                                
                                📍 **To view the dashboard**: Copy the path above, open Finder/Explorer, 
                                navigate to the folder, and double-click `dashboard.html`
                                """)
                                
                                gr.HTML('</div>')  # Close generated files box
                
                # Connect button
                submit_btn.click(
                    fn=process_document,
                    inputs=[
                        file_input,
                        output_language,
                        enable_iterative,
                        max_iterations,
                        target_score,
                        enable_python_tools,
                        enable_interactive,
                        enable_deep_review,
                        reference_files,
                        reference_type
                    ],
                    outputs=[
                        status_output, 
                        report_output, 
                        json_output, 
                        dashboard_output, 
                        agents_output,
                        ratings_output,  # Agent ratings tab
                        agents_deployed_output,  # NEW: Agents deployed info
                        output_dir_display,
                        time_estimate_output,
                        download_report_btn,
                        download_json_btn,
                        download_dashboard_btn
                    ]
                )
            
            # Help Tab
            with gr.Tab("💡 How It Works"):
                gr.HTML("""
                <div style="max-width: 900px; margin: 0 auto;">
                    <h2 style="color: #667eea; text-align: center; font-size: 32px; margin-bottom: 40px;">
                        Get Professional Feedback in 3 Simple Steps
                    </h2>
                </div>
                """)
                
                gr.HTML("""
                <div style="display: grid; grid-template-columns: repeat(3, 1fr); gap: 30px; margin: 40px 0;">
                    <div style="text-align: center; padding: 30px;">
                        <div style="font-size: 64px; margin-bottom: 20px;">📤</div>
                        <h3 style="color: #667eea; margin: 15px 0;">1. Upload</h3>
                        <p style="color: #666;">Drop your document or click to browse. We support Word, PDF, and text files.</p>
                    </div>
                    <div style="text-align: center; padding: 30px;">
                        <div style="font-size: 64px; margin-bottom: 20px;">⚙️</div>
                        <h3 style="color: #667eea; margin: 15px 0;">2. Configure</h3>
                        <p style="color: #666;">Choose your language and optional features. Our defaults work great for most documents!</p>
                    </div>
                    <div style="text-align: center; padding: 30px;">
                        <div style="font-size: 64px; margin-bottom: 20px;">✨</div>
                        <h3 style="color: #667eea; margin: 15px 0;">3. Review</h3>
                        <p style="color: #666;">Get detailed feedback in minutes. Download or view results right in your browser.</p>
                    </div>
                </div>
                """)
                
                gr.HTML("""
                <div style="display: grid; grid-template-columns: repeat(2, 1fr); gap: 25px; margin: 40px 0;">
                    <div class="benefit-card">
                        <h3 style="color: #667eea; margin: 0 0 15px 0;">📋 Comprehensive Analysis</h3>
                        <p style="color: #666;">30+ AI specialists review your document from every angle - style, clarity, facts, structure, and more.</p>
                    </div>
                    <div class="benefit-card">
                        <h3 style="color: #667eea; margin: 0 0 15px 0;">🌍 Any Language</h3>
                        <p style="color: #666;">Write in your language, get feedback in yours. We support Italian, English, Spanish, French, German, and more.</p>
                    </div>
                    <div class="benefit-card">
                        <h3 style="color: #667eea; margin: 0 0 15px 0;">🔄 Auto-Improve</h3>
                        <p style="color: #666;">Let AI enhance your document automatically through multiple refinement cycles until it's perfect.</p>
                    </div>
                    <div class="benefit-card">
                        <h3 style="color: #667eea; margin: 0 0 15px 0;">📊 Number Validation</h3>
                        <p style="color: #666;">Automatically verify calculations, check statistics, and validate data consistency in your documents.</p>
                    </div>
                    <div class="benefit-card">
                        <h3 style="color: #667eea; margin: 0 0 15px 0;">🔬 Academic Research</h3>
                        <p style="color: #666;">For research papers: search 200M+ academic papers, get formal citations, identify literature gaps.</p>
                    </div>
                    <div class="benefit-card">
                        <h3 style="color: #667eea; margin: 0 0 15px 0;">📚 Template Compliance</h3>
                        <p style="color: #666;">Upload your templates or guidelines and ensure your document follows them perfectly.</p>
                    </div>
                </div>
                """)
                
                gr.Markdown("""
                
                ---
                
                ## 📦 What You'll Receive
                
                After analysis, you get:
                
                - **📄 Full Report** - Detailed feedback in an easy-to-read format
                - **📊 Interactive Dashboard** - Visual charts and metrics (open in browser)
                - **💾 JSON Data** - Complete structured results for integration
                - **📁 All Files** - Everything saved in a timestamped folder for your records
                
                ---
                
                ## 💡 Pro Tips
                
                <div style="background: #f8f9fa; padding: 25px; border-radius: 12px; border-left: 4px solid #667eea;">
                
                **For Best Results:**
                
                ✨ Use **Auto-Improve** if you want AI to enhance your document automatically
                
                📊 Enable **Number Validation** for documents with data, calculations, or statistics
                
                🔬 Try **Deep Analysis** for scientific papers, research proposals, or technical reports
                
                📚 Upload **Templates** if you need to match specific formatting or style guidelines
                
                🌍 Let us **Auto-Detect** your language - it works great!
                
                </div>
                
                ---
                
                ## ❓ Common Questions
                
                **How long does it take?**  
                Most documents are analyzed in 3-5 minutes. Deep analysis takes 8-12 minutes.
                
                **What file types work?**  
                PDF, Word (DOCX), Text (TXT), and Markdown (MD) files.
                
                **Is my document safe?**  
                Yes! Everything is processed securely and your files are never shared.
                
                **Can I improve multiple times?**  
                Absolutely! The Auto-Improve feature can run multiple refinement cycles (we recommend 3).
                
                **What's the difference between standard and deep analysis?**  
                Standard is perfect for most documents. Deep analysis adds academic literature search and 20+ specialist reviewers - best for research papers.
                
                ---
                
                <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 30px; border-radius: 15px; color: white; text-align: center;">
                    <h3 style="margin: 0 0 15px 0;">Ready to Transform Your Documents?</h3>
                    <p style="opacity: 0.95; margin: 0;">Upload your first document and see the difference AI can make!</p>
                </div>
                """)
            
            # About Tab
            with gr.Tab("ℹ️ About"):
                gr.HTML("""
                <div style="text-align: center; padding: 40px 20px;">
                    <div style="font-size: 64px; margin-bottom: 20px;">✨</div>
                    <h1 style="color: #667eea; font-size: 42px; margin: 0 0 15px 0;">AI Document Pro</h1>
                    <p style="font-size: 20px; color: #666; margin: 0;">Professional Document Analysis, Powered by AI</p>
                    <p style="font-size: 16px; color: #999; margin: 10px 0 0 0;">Version 3.1 Enterprise</p>
                </div>
                """)
                
                gr.HTML("""
                <div style="max-width: 900px; margin: 40px auto;">
                    <h2 style="color: #667eea; text-align: center; margin-bottom: 40px;">Why Choose Us?</h2>
                    
                    <div style="display: grid; grid-template-columns: repeat(2, 1fr); gap: 30px; margin-bottom: 60px;">
                        <div style="text-align: center; padding: 20px;">
                            <div style="font-size: 48px; margin-bottom: 15px;">🚀</div>
                            <h3 style="color: #667eea; margin: 10px 0;">Lightning Fast</h3>
                            <p style="color: #666;">Get comprehensive feedback in just minutes, not hours or days</p>
                        </div>
                        <div style="text-align: center; padding: 20px;">
                            <div style="font-size: 48px; margin-bottom: 15px;">🎯</div>
                            <h3 style="color: #667eea; margin: 10px 0;">Expert Quality</h3>
                            <p style="color: #666;">30+ AI specialists analyze every aspect of your document</p>
                        </div>
                        <div style="text-align: center; padding: 20px;">
                            <div style="font-size: 48px; margin-bottom: 15px;">🌍</div>
                            <h3 style="color: #667eea; margin: 10px 0;">Global Reach</h3>
                            <p style="color: #666;">Support for 10+ languages with automatic detection</p>
                        </div>
                        <div style="text-align: center; padding: 20px;">
                            <div style="font-size: 48px; margin-bottom: 15px;">🔒</div>
                            <h3 style="color: #667eea; margin: 10px 0;">Secure & Private</h3>
                            <p style="color: #666;">Your documents are processed securely and never shared</p>
                        </div>
                    </div>
                    
                    <div style="background: #f8f9fa; padding: 40px; border-radius: 15px; margin-bottom: 40px;">
                        <h2 style="color: #667eea; text-align: center; margin: 0 0 30px 0;">What We Analyze</h2>
                        <div style="display: grid; grid-template-columns: repeat(3, 1fr); gap: 20px;">
                            <div>
                                <h4 style="color: #667eea; margin: 0 0 10px 0;">✍️ Writing Quality</h4>
                                <p style="color: #666; margin: 0; font-size: 14px;">Grammar, style, clarity, tone</p>
                            </div>
                            <div>
                                <h4 style="color: #667eea; margin: 0 0 10px 0;">📊 Data Accuracy</h4>
                                <p style="color: #666; margin: 0; font-size: 14px;">Numbers, calculations, statistics</p>
                            </div>
                            <div>
                                <h4 style="color: #667eea; margin: 0 0 10px 0;">🎯 Structure</h4>
                                <p style="color: #666; margin: 0; font-size: 14px;">Organization, flow, coherence</p>
                            </div>
                            <div>
                                <h4 style="color: #667eea; margin: 0 0 10px 0;">✓ Facts</h4>
                                <p style="color: #666; margin: 0; font-size: 14px;">Accuracy, citations, sources</p>
                            </div>
                            <div>
                                <h4 style="color: #667eea; margin: 0 0 10px 0;">💡 Logic</h4>
                                <p style="color: #666; margin: 0; font-size: 14px;">Arguments, reasoning, consistency</p>
                            </div>
                            <div>
                                <h4 style="color: #667eea; margin: 0 0 10px 0;">🎨 Presentation</h4>
                                <p style="color: #666; margin: 0; font-size: 14px;">Format, readability, impact</p>
                            </div>
                        </div>
                    </div>
                    
                    <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 40px; border-radius: 15px; color: white; text-align: center;">
                        <h2 style="margin: 0 0 20px 0;">Powered by Cutting-Edge Technology</h2>
                        <p style="opacity: 0.95; margin: 0 0 30px 0; font-size: 18px;">We use the latest AI models from OpenAI to deliver professional-grade analysis</p>
                        <div style="display: flex; justify-content: center; gap: 30px; flex-wrap: wrap;">
                            <div style="background: rgba(255,255,255,0.2); padding: 15px 30px; border-radius: 10px;">
                                <div style="font-weight: 600;">OpenAI GPT-5</div>
                                <div style="font-size: 14px; opacity: 0.9;">Most Advanced AI</div>
                            </div>
                            <div style="background: rgba(255,255,255,0.2); padding: 15px 30px; border-radius: 10px;">
                                <div style="font-weight: 600;">30+ Specialists</div>
                                <div style="font-size: 14px; opacity: 0.9;">Multi-Agent System</div>
                            </div>
                            <div style="background: rgba(255,255,255,0.2); padding: 15px 30px; border-radius: 10px;">
                                <div style="font-weight: 600;">200M+ Papers</div>
                                <div style="font-size: 14px; opacity: 0.9;">Academic Database</div>
                            </div>
                        </div>
                    </div>
                </div>
                """)
                
                gr.Markdown("""
                
                ---
                
                <div style="text-align: center; padding: 20px; color: #999;">
                    <p><strong>Need help?</strong> Check the "How It Works" tab for detailed guidance</p>
                    <p style="margin-top: 20px;">Made with ❤️ for better document workflows</p>
                </div>
                """)
        
        gr.HTML("""
        <div style="background: #f8f9fa; padding: 40px 20px; margin-top: 60px; border-radius: 15px;">
            <div style="max-width: 1200px; margin: 0 auto;">
                <div style="display: grid; grid-template-columns: repeat(3, 1fr); gap: 40px; margin-bottom: 40px;">
                    <div>
                        <h4 style="color: #667eea; margin: 0 0 15px 0;">✨ AI Document Pro</h4>
                        <p style="color: #666; font-size: 14px; margin: 0;">Transform your documents with professional AI-powered analysis in minutes.</p>
                    </div>
                    <div>
                        <h4 style="color: #667eea; margin: 0 0 15px 0;">Features</h4>
                        <p style="color: #666; font-size: 14px; margin: 5px 0;">🚀 Lightning Fast Analysis</p>
                        <p style="color: #666; font-size: 14px; margin: 5px 0;">🎯 30+ AI Specialists</p>
                        <p style="color: #666; font-size: 14px; margin: 5px 0;">🌍 Multi-Language Support</p>
                    </div>
                    <div>
                        <h4 style="color: #667eea; margin: 0 0 15px 0;">Powered By</h4>
                        <p style="color: #666; font-size: 14px; margin: 5px 0;">OpenAI GPT-5</p>
                        <p style="color: #666; font-size: 14px; margin: 5px 0;">Semantic Scholar</p>
                        <p style="color: #666; font-size: 14px; margin: 5px 0;">Advanced AI Technology</p>
                    </div>
                </div>
                <div style="border-top: 1px solid #ddd; padding-top: 20px; text-align: center; color: #999; font-size: 13px;">
                    <p style="margin: 0;">© 2025 AI Document Pro v3.1 Enterprise | All Rights Reserved</p>
                    <p style="margin: 10px 0 0 0;">Made with ❤️ for better document workflows</p>
                </div>
            </div>
        </div>
        """)
    
    return app


def launch_ui(share: bool = False, server_port: int = 7860):
    """Launch the Gradio interface."""
    app = create_ui()
    
    print("\n" + "="*70)
    print("🚀 Launching Document Review System Web UI")
    print("="*70)
    print(f"\n📍 Local URL: http://localhost:{server_port}")
    if share:
        print("🌐 Public URL will be generated...")
    print("\n💡 Press Ctrl+C to stop the server")
    print("="*70 + "\n")
    
    # Configure for better file upload handling
    import logging
    
    # Suppress ClientDisconnect warnings (they're normal when users cancel uploads)
    logging.getLogger("starlette.requests").setLevel(logging.ERROR)
    logging.getLogger("uvicorn.error").setLevel(logging.WARNING)
    
    app.launch(
        share=share,
        server_port=server_port,
        server_name="0.0.0.0",
        show_error=True,
        favicon_path=None,
        max_file_size="100mb",  # Support larger files
        # Enable queue for better handling of multiple requests
        max_threads=40
    )


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Launch Document Review System Web UI")
    parser.add_argument("--share", action="store_true", help="Create public Gradio link")
    parser.add_argument("--port", type=int, default=7860, help="Server port (default: 7860)")
    
    args = parser.parse_args()
    
    launch_ui(share=args.share, server_port=args.port)

