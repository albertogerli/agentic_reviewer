"""
Reference Context System
Loads reference documents as context for document development/review.

Use cases:
- Templates as reference for new documents
- Company guidelines for proposals
- Previous successful documents as examples
- Style guides for writing
- Data sheets for technical documents
"""

import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import json

logger = logging.getLogger(__name__)


@dataclass
class ReferenceDocument:
    """A reference document providing context."""
    file_path: str
    file_name: str
    content: str
    document_type: str  # template, guideline, example, data, style_guide
    description: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None


class ReferenceContextManager:
    """
    Manages reference documents that provide context for document development.
    
    Reference documents can be:
    - Templates (structure/format examples)
    - Guidelines (company policies, writing standards)
    - Examples (previous successful documents)
    - Data sheets (technical specifications, product data)
    - Style guides (tone, formatting rules)
    """
    
    def __init__(self):
        self.references: List[ReferenceDocument] = []
    
    def load_references(self, input_path: str, 
                       document_type: str = "example",
                       description: Optional[str] = None) -> int:
        """
        Load reference documents from path (file, directory, or ZIP).
        
        Args:
            input_path: Path to file/directory/ZIP
            document_type: Type of reference (template, guideline, example, data, style_guide)
            description: Optional description of references
        
        Returns:
            Number of references loaded
        """
        path = Path(input_path)
        
        if not path.exists():
            logger.warning(f"Reference path not found: {input_path}")
            return 0
        
        loaded = 0
        
        if path.is_file() and path.suffix.lower() == '.zip':
            loaded = self._load_from_zip(path, document_type, description)
        elif path.is_dir():
            loaded = self._load_from_directory(path, document_type, description)
        elif path.is_file():
            ref = self._load_single_file(path, document_type, description)
            if ref:
                self.references.append(ref)
                loaded = 1
        
        logger.info(f"Loaded {loaded} reference documents from {input_path}")
        return loaded
    
    def _load_from_directory(self, directory: Path, 
                            document_type: str,
                            description: Optional[str]) -> int:
        """Load all supported files from directory."""
        loaded = 0
        supported_exts = {'.pdf', '.txt', '.md', '.docx', '.doc', '.json'}
        
        for file_path in directory.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in supported_exts:
                ref = self._load_single_file(file_path, document_type, description)
                if ref:
                    self.references.append(ref)
                    loaded += 1
        
        return loaded
    
    def _load_from_zip(self, zip_path: Path, 
                      document_type: str,
                      description: Optional[str]) -> int:
        """Load files from ZIP archive."""
        import zipfile
        import tempfile
        
        loaded = 0
        
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                    zip_ref.extractall(temp_dir)
                
                loaded = self._load_from_directory(
                    Path(temp_dir), document_type, description
                )
        except Exception as e:
            logger.error(f"Error loading ZIP: {e}")
        
        return loaded
    
    def _load_single_file(self, file_path: Path, 
                         document_type: str,
                         description: Optional[str]) -> Optional[ReferenceDocument]:
        """Load a single reference file."""
        try:
            content = self._read_file_content(file_path)
            
            if not content:
                return None
            
            return ReferenceDocument(
                file_path=str(file_path),
                file_name=file_path.name,
                content=content,
                document_type=document_type,
                description=description or f"{document_type.capitalize()}: {file_path.name}"
            )
            
        except Exception as e:
            logger.error(f"Error loading {file_path}: {e}")
            return None
    
    def _read_file_content(self, file_path: Path) -> str:
        """Read content from file."""
        try:
            if file_path.suffix.lower() == '.pdf':
                return self._read_pdf(file_path)
            elif file_path.suffix.lower() in ['.txt', '.md', '.json']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                return self._read_word(file_path)
            else:
                return ""
        except Exception as e:
            logger.error(f"Error reading {file_path}: {e}")
            return ""
    
    def _read_pdf(self, file_path: Path) -> str:
        """Read PDF content."""
        try:
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = []
                for page in reader.pages:
                    text.append(page.extract_text())
                return '\n'.join(text)
        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            return ""
    
    def _read_word(self, file_path: Path) -> str:
        """Read Word document content."""
        try:
            import docx
            doc = docx.Document(file_path)
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        except ImportError:
            logger.warning("python-docx not installed")
            return ""
        except Exception as e:
            logger.error(f"Error reading Word: {e}")
            return ""
    
    def get_context_for_review(self, include_types: Optional[List[str]] = None,
                               max_total_chars: int = 50000) -> str:
        """
        Get formatted context from reference documents for AI review.
        
        Args:
            include_types: Types to include (None = all)
            max_total_chars: Maximum total characters to include
        
        Returns:
            Formatted context string
        """
        if not self.references:
            return ""
        
        # Filter by type if specified
        refs_to_include = self.references
        if include_types:
            refs_to_include = [r for r in self.references if r.document_type in include_types]
        
        if not refs_to_include:
            return ""
        
        context_parts = []
        context_parts.append("# REFERENCE DOCUMENTS FOR CONTEXT\n")
        context_parts.append("The following reference documents provide context for this review:\n")
        
        total_chars = 0
        
        for ref in refs_to_include:
            # Add header
            header = f"\n## {ref.description or ref.file_name}\n"
            header += f"**Type:** {ref.document_type}\n"
            header += f"**Source:** {ref.file_name}\n\n"
            
            # Calculate remaining space
            remaining = max_total_chars - total_chars - len(header)
            
            if remaining <= 100:
                # Not enough space
                context_parts.append("\n... (additional references omitted due to length)\n")
                break
            
            # Add content (truncated if needed)
            content = ref.content
            if len(content) > remaining:
                content = content[:remaining-50] + "\n\n... (content truncated)"
            
            context_parts.append(header)
            context_parts.append(content)
            context_parts.append("\n\n---\n")
            
            total_chars += len(header) + len(content)
        
        return ''.join(context_parts)
    
    def get_summary(self) -> Dict[str, Any]:
        """Get summary of loaded references."""
        by_type = {}
        for ref in self.references:
            by_type[ref.document_type] = by_type.get(ref.document_type, 0) + 1
        
        total_size = sum(len(ref.content) for ref in self.references)
        
        return {
            'total_references': len(self.references),
            'by_type': by_type,
            'total_content_size': total_size,
            'references': [
                {
                    'file_name': ref.file_name,
                    'type': ref.document_type,
                    'size': len(ref.content)
                }
                for ref in self.references
            ]
        }
    
    def clear(self):
        """Clear all loaded references."""
        self.references = []
        logger.info("Cleared all reference documents")


class ReferenceAugmentedPrompt:
    """
    Augments AI prompts with reference document context.
    """
    
    @staticmethod
    def create_review_prompt_with_context(document_text: str,
                                         reference_context: str,
                                         agent_instructions: str,
                                         output_language: str = "English") -> str:
        """
        Create review prompt that includes reference context.
        
        Args:
            document_text: The document to review
            reference_context: Context from reference documents
            agent_instructions: Agent's specific instructions
            output_language: Output language
        
        Returns:
            Complete prompt with context
        """
        prompt = f"""{agent_instructions}

**OUTPUT LANGUAGE:** {output_language}

{reference_context if reference_context else ""}

# DOCUMENT TO REVIEW

Please review the following document. Consider the reference documents above as context for expectations, standards, and examples.

{document_text}

---

Provide your review considering:
1. How well it matches reference templates/examples (if provided)
2. Compliance with guidelines (if provided)
3. Comparison to successful examples (if provided)
4. Adherence to style guides (if provided)
5. Accuracy against reference data (if provided)

**IMPORTANT:** End your review with "REVIEW COMPLETED - [Your Agent Name]"
"""
        
        return prompt
    
    @staticmethod
    def create_improvement_prompt_with_context(document_text: str,
                                              reviews: str,
                                              reference_context: str) -> str:
        """
        Create improvement prompt with reference context.
        """
        prompt = f"""You are an expert document editor. Improve the document based on reviews and reference context.

{reference_context if reference_context else ""}

CURRENT DOCUMENT:
{document_text}

EXPERT REVIEWS AND FEEDBACK:
{reviews}

Your task:
1. Apply suggested improvements
2. Align with reference templates/standards (if provided)
3. Match style of successful examples (if provided)
4. Follow guidelines strictly (if provided)
5. Maintain accuracy with reference data (if provided)

Return the improved document.
"""
        
        return prompt


def create_example_reference_config() -> Dict[str, Any]:
    """
    Create example configuration for reference documents.
    
    Returns example JSON structure users can adapt.
    """
    return {
        "reference_documents": {
            "templates": {
                "path": "/path/to/templates/",
                "type": "template",
                "description": "Company document templates"
            },
            "guidelines": {
                "path": "/path/to/guidelines.pdf",
                "type": "guideline",
                "description": "Corporate writing guidelines"
            },
            "examples": {
                "path": "/path/to/successful_proposals.zip",
                "type": "example",
                "description": "Successful past proposals"
            },
            "style_guides": {
                "path": "/path/to/style_guide.md",
                "type": "style_guide",
                "description": "Company style guide"
            },
            "data_sheets": {
                "path": "/path/to/product_specs/",
                "type": "data",
                "description": "Product technical specifications"
            }
        }
    }

