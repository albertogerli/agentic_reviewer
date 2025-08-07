from __future__ import annotations

import io
import json
import mimetypes
import os
from dataclasses import dataclass, field
from datetime import datetime
from typing import Any, Dict, Optional, Tuple, Union

# Optional imports guarded at runtime
try:
    import chardet  # type: ignore
except Exception:  # pragma: no cover - optional
    chardet = None  # type: ignore

try:
    from pdfminer.high_level import extract_text as pdf_extract_text  # type: ignore
except Exception:  # pragma: no cover - optional
    pdf_extract_text = None  # type: ignore

try:
    import docx  # python-docx  # type: ignore
except Exception:  # pragma: no cover - optional
    docx = None  # type: ignore

try:
    from bs4 import BeautifulSoup  # type: ignore
except Exception:  # pragma: no cover - optional
    BeautifulSoup = None  # type: ignore

try:
    import textract  # type: ignore
except Exception:  # pragma: no cover - optional
    textract = None  # type: ignore


@dataclass
class Document:
    text: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    source_path: Optional[str] = None

    def to_dict(self) -> Dict[str, Any]:
        return {
            "text": self.text,
            "metadata": self.metadata,
            "source_path": self.source_path,
        }


def _detect_mime_and_encoding(
    *,
    path: Optional[str],
    raw: Optional[bytes],
) -> Tuple[Optional[str], Optional[str]]:
    """
    Best-effort MIME and encoding detection.

    - Uses mimetypes by extension for path-based MIME.
    - For encoding, tries chardet if available, otherwise UTF-8 fallback.
    - For raw bytes only, attempts to infer MIME for common signatures (PDF, HTML, plain text).
    """
    mime: Optional[str] = None
    encoding: Optional[str] = None

    if path:
        mime, _ = mimetypes.guess_type(path)

    sample = None
    if raw is not None:
        sample = raw[:4096]
    elif path and os.path.exists(path) and os.path.isfile(path):
        try:
            with open(path, "rb") as f:
                sample = f.read(4096)
        except Exception:
            sample = None

    # Heuristic MIME for some common formats when not known
    if mime is None and sample is not None:
        if sample.startswith(b"%PDF"):
            mime = "application/pdf"
        elif sample.lstrip().startswith(b"<") and b"html" in sample[:1024].lower():
            mime = "text/html"
        elif b"{" in sample[:1] or b"[" in sample[:1]:
            # naive hint for JSON if starts with { or [
            try:
                json.loads(sample.decode("utf-8", errors="ignore"))
                mime = "application/json"
            except Exception:
                pass
        else:
            # default text if it mostly decodes
            try:
                sample.decode("utf-8")
                mime = mime or "text/plain"
            except Exception:
                pass

    # Encoding detection
    if sample is not None:
        if chardet is not None:
            try:
                guess = chardet.detect(sample)
                encoding = (guess or {}).get("encoding")
            except Exception:
                encoding = None
        if not encoding:
            # fallback: assume utf-8 unless it raises immediately
            try:
                sample.decode("utf-8")
                encoding = "utf-8"
            except Exception:
                encoding = None

    return mime, encoding


def _read_all_bytes(path: str) -> bytes:
    with open(path, "rb") as f:
        return f.read()


def _extract_text_from_pdf(path: Optional[str] = None, raw: Optional[bytes] = None) -> str:
    if pdf_extract_text is None:
        raise RuntimeError("pdfminer.six is required to extract text from PDFs")
    if path:
        return pdf_extract_text(path) or ""
    assert raw is not None
    # Write to a temporary buffer file-like object is not directly supported by pdfminer.high_level
    # So for raw bytes, write to a temporary NamedTemporaryFile
    import tempfile

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=True) as tmp:
        tmp.write(raw)
        tmp.flush()
        return pdf_extract_text(tmp.name) or ""


def _extract_text_from_docx(path: Optional[str] = None, raw: Optional[bytes] = None) -> str:
    if docx is None:
        raise RuntimeError("python-docx is required to extract text from DOCX files")
    if path:
        document = docx.Document(path)
    else:
        assert raw is not None
        document = docx.Document(io.BytesIO(raw))
    paragraphs = [p.text for p in document.paragraphs]
    return "\n".join([t for t in paragraphs if t is not None])


def _extract_text_from_html(raw_text: str) -> str:
    if BeautifulSoup is None:
        # fallback: return text as-is
        return raw_text
    soup = BeautifulSoup(raw_text, "html.parser")
    # remove script/style
    for tag in soup(["script", "style"]):
        tag.decompose()
    return soup.get_text("\n")


def _extract_text_via_textract(path: Optional[str] = None, raw: Optional[bytes] = None) -> str:
    if textract is None:
        raise RuntimeError("textract is not installed")
    if path:
        data = textract.process(path)
    else:
        assert raw is not None
        # textract prefers a filename; use temporary file for arbitrary bytes
        import tempfile

        with tempfile.NamedTemporaryFile(delete=True) as tmp:
            tmp.write(raw)
            tmp.flush()
            data = textract.process(tmp.name)
    try:
        return data.decode("utf-8")
    except Exception:
        return data.decode(errors="ignore")


def _normalize_text(text: str) -> str:
    # Simple normalization: strip, normalize newlines, collapse excessive whitespace
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # collapse 3+ newlines to 2
    import re

    text = re.sub(r"\n{3,}", "\n\n", text)
    # strip trailing spaces on lines
    text = "\n".join([line.rstrip() for line in text.splitlines()])
    return text.strip()


def ingest(
    *,
    path: Optional[str] = None,
    raw: Optional[Union[bytes, bytearray]] = None,
    filename: Optional[str] = None,
    extra_metadata: Optional[Dict[str, Any]] = None,
) -> Document:
    """
    Ingest a document from a file path or raw bytes, detect MIME and encoding, extract text,
    and return a normalized Document dataclass.

    Args:
        path: Optional filesystem path to a file.
        raw: Optional bytes of the file content.
        filename: Optional hint for filename when using raw bytes (affects MIME guess by extension).
        extra_metadata: Optional extra metadata to attach to the document.

    Returns:
        Document: containing extracted text, metadata, and original source path if provided.

    Raises:
        ValueError: if neither path nor raw is provided.
    """
    if path is None and raw is None:
        raise ValueError("Provide either 'path' or 'raw'")

    if isinstance(raw, bytearray):
        raw = bytes(raw)

    # Prefer actual path for extension-based mime; else use provided filename
    path_for_mime = path or filename

    mime, encoding = _detect_mime_and_encoding(path=path_for_mime, raw=raw)

    extracted_text = ""

    # Branch by MIME or extension hints
    lower_name = (path_for_mime or "").lower()

    def decode_raw_with(enc: Optional[str]) -> Optional[str]:
        if raw is None:
            return None
        encs = [e for e in [enc, "utf-8", "latin-1"] if e]
        for e in encs:
            try:
                return raw.decode(e)
            except Exception:
                continue
        try:
            return raw.decode(errors="ignore")
        except Exception:
            return None

    try:
        if (mime == "application/pdf") or lower_name.endswith(".pdf"):
            extracted_text = _extract_text_from_pdf(path=path, raw=raw)
        elif (mime in {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}) or lower_name.endswith(".docx"):
            extracted_text = _extract_text_from_docx(path=path, raw=raw)
        elif (mime in {"text/html"}) or lower_name.endswith((".html", ".htm")):
            text_raw = None
            if path and raw is None:
                with open(path, "rb") as f:
                    raw_bytes = f.read()
                text_raw = decode_raw_with(encoding) or raw_bytes.decode("utf-8", errors="ignore")
            else:
                text_raw = decode_raw_with(encoding) or ""
            extracted_text = _extract_text_from_html(text_raw)
        elif (mime in {"application/json"}) or lower_name.endswith(".json"):
            text_raw = None
            if path and raw is None:
                with open(path, "rb") as f:
                    raw_bytes = f.read()
                text_raw = decode_raw_with(encoding) or raw_bytes.decode("utf-8", errors="ignore")
            else:
                text_raw = decode_raw_with(encoding) or ""
            try:
                parsed = json.loads(text_raw)
                extracted_text = json.dumps(parsed, ensure_ascii=False, indent=2)
            except Exception:
                extracted_text = text_raw or ""
        elif mime and mime.startswith("text/") or lower_name.endswith(
            (".txt", ".md", ".csv", ".tsv", ".py", ".js", ".jsonl")
        ):
            if path and raw is None:
                with open(path, "rb") as f:
                    raw_bytes = f.read()
                extracted_text = decode_raw_with(encoding) or raw_bytes.decode("utf-8", errors="ignore")
            else:
                extracted_text = decode_raw_with(encoding) or ""
        else:
            # Fallback: use textract if present; otherwise try to decode as text
            if textract is not None:
                extracted_text = _extract_text_via_textract(path=path, raw=raw)
            else:
                extracted_text = decode_raw_with(encoding) or ""
    except Exception as e:
        # As final fallback, try textract if not already tried
        if textract is not None and not extracted_text:
            try:
                extracted_text = _extract_text_via_textract(path=path, raw=raw)
            except Exception:
                raise e
        else:
            raise e

    extracted_text = _normalize_text(extracted_text)

    stat_info = None
    try:
        if path and os.path.exists(path):
            st = os.stat(path)
            stat_info = {
                "size": st.st_size,
                "modified": datetime.fromtimestamp(st.st_mtime).isoformat(),
                "created": datetime.fromtimestamp(st.st_ctime).isoformat(),
            }
    except Exception:
        stat_info = None

    metadata: Dict[str, Any] = {
        "mime": mime,
        "encoding": encoding,
        "filename": os.path.basename(path) if path else (os.path.basename(filename) if filename else None),
    }
    if stat_info:
        metadata["file_stat"] = stat_info
    if extra_metadata:
        metadata.update(extra_metadata)

    return Document(text=extracted_text, metadata=metadata, source_path=path)

